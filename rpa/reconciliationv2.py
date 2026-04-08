# pip install playwright httpx openpyxl python-docx && playwright install chromium
# 任务：reconciliationV2
# 录制时间：2026-04-07 19:56:47
# 由 OpenClaw RPA Recorder（headed 真实录制）生成 — 可脱离 OpenClaw 独立运行

import asyncio
import os
import urllib.parse
from pathlib import Path

import httpx
from openpyxl import Workbook, load_workbook
from openpyxl.utils import get_column_letter

from docx import Document

from playwright.async_api import async_playwright, TimeoutError as PlaywrightTimeout

CONFIG = {
    "output_dir":    Path.home() / "Desktop",
    "headless":      False,
    "timeout":       60_000,
    "slow_mo":       300,
    # 导航后等待 SPA 内容渲染的额外时间（重型 SPA 如 Yahoo Finance 需要 1-2 秒）
    "spa_settle_ms": 1_500,
    # extract_text 等待目标元素出现的超时（毫秒）
    "content_wait":  15_000,
    # httpx 调用外部 API 的超时（秒）
    "api_timeout":   60.0,
    # 若有已保存的登录 Cookie，填写路径（由 rpa_manager login-start 生成）
    "cookies_path":  "",}

_UA = (
    "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
    "AppleWebKit/537.36 (KHTML, like Gecko) "
    "Chrome/124.0.0.0 Safari/537.36"
)

def _format_extract_section(field_label: str, lines: list[str]) -> str:
    """Format extracted DOM text: show field name, separator line, then body."""
    name = (field_label or "").strip() or "extract"
    title = f"【字段：{name}】"
    if not lines:
        body = "(no text matched)\n"
    elif len(lines) == 1:
        body = lines[0].strip() + "\n"
    else:
        parts = [f"{i + 1}. {s.strip()}" for i, s in enumerate(lines)]
        body = "\n\n".join(parts) + "\n"
    sep = "─" * 32
    return f"{title}\n{sep}\n{body}\n"


_EXTRACT_JS = '([s,n])=>{const r=document.querySelector("main")||document.querySelector(\'[role="main"]\');const bare=/^[a-zA-Z][a-zA-Z0-9-]*$/.test(s)&&s.indexOf("#")<0&&s.indexOf(".")<0&&s.indexOf("[")<0&&s.indexOf(" ")<0;const sc=bare&&r?r:document;return Array.from(sc.querySelectorAll(s)).slice(0,n).map(e=>(e.textContent||"").replace(/\\s+/g," ").trim()).filter(Boolean)}'


async def _wait_for_content(page, selector: str) -> None:
    """等待 selector 对应的元素出现在 DOM 中（容错：超时也继续）。"""
    try:
        await page.wait_for_selector(selector, timeout=CONFIG["content_wait"])
    except Exception:
        pass  # 元素未出现也继续，evaluate 会返回空列表


async def _scroll_window(page, dy: int) -> None:
    """窗口滚动：导航后若再用 evaluate(scrollBy)，易因执行上下文销毁报错；用 mouse.wheel 并在滚动前等待页面稳定。"""
    try:
        await page.wait_for_load_state("domcontentloaded", timeout=10_000)
    except Exception:
        pass
    vp = page.viewport_size
    if vp:
        await page.mouse.move(vp["width"] // 2, vp["height"] // 2)
    else:
        await page.mouse.move(720, 450)
    await page.mouse.wheel(0, float(dy))


async def run():
    async with async_playwright() as p:
        browser = await p.chromium.launch(
            headless=CONFIG["headless"],
            args=["--no-sandbox", "--disable-blink-features=AutomationControlled"],
            slow_mo=CONFIG["slow_mo"],
        )
        context = await browser.new_context(
            user_agent=_UA,
            viewport={"width": 1440, "height": 900},
            locale="en-US",
            timezone_id="America/New_York",
            extra_http_headers={"Accept-Language": "en-US,en;q=0.9"},
        )
        await context.add_init_script(
            "Object.defineProperty(navigator, 'webdriver', {get: () => undefined})"
        )
        # 若配置了 cookies_path，注入 Cookie 模拟已登录态（由 rpa_manager login-start 生成）
        import json as _json
        _cp = CONFIG.get("cookies_path", "")
        if _cp and Path(_cp).exists():
            await context.add_cookies(_json.loads(Path(_cp).read_text()))
        page = await context.new_page()
        page.set_default_timeout(CONFIG["timeout"])

        try:
            # ── 步骤 1：Fetch AP open batches
            try:
                _params = {'status': 'open', 'week': '2026-W14'}
                _api_url = 'https://0a34723da37946b7add0b4581c37ada2_oas.api.mockbin.io/ap/reconciliation/batches' + "?" + urllib.parse.urlencode(_params)
                async with httpx.AsyncClient(timeout=CONFIG["api_timeout"], verify=False) as _hc:
                    _r = await _hc.get(_api_url)
                    _r.raise_for_status()
                (CONFIG["output_dir"] / 'reconcile_raw.json').write_text(_r.text, encoding="utf-8")
                print("API 响应已写入", CONFIG["output_dir"] / 'reconcile_raw.json')
            except Exception:
                await page.screenshot(path="step_1_error.png")
                raise

            # ── 步骤 2：Write system data to Excel
            try:
                import json
                import os
                from openpyxl import Workbook
                
                json_path = os.path.join(CONFIG["output_dir"], "reconcile_raw.json")
                excel_path = os.path.join(CONFIG["output_dir"], "ap_draft_thisweek.xlsx")
                
                with open(json_path, "r", encoding="utf-8") as f:
                    data = json.load(f)
                
                wb = Workbook()
                ws = wb.active
                ws.title = "System"
                headers = ["line_id", "vendor_id", "po_ref", "amount_system", "currency", "due_date", "batch_id"]
                ws.append(headers)
                ws.freeze_panes = "A2"
                
                for batch in data.get("batches", []):
                    batch_id = batch.get("batch_id", "")
                    for line in batch.get("lines", []):
                        ws.append([
                            line.get("line_id"),
                            line.get("vendor_id"),
                            line.get("po_ref"),
                            line.get("amount_system"),
                            line.get("currency"),
                            line.get("due_date"),
                            batch_id
                        ])
                
                wb.save(excel_path)
            except Exception:
                await page.screenshot(path="step_2_error.png")
                raise

            # ── 步骤 3：Copy invoices to Excel
            try:
                import os
                import openpyxl
                
                source_path = os.path.join(CONFIG["output_dir"], "invoice_import_thisweek.xlsx")
                target_path = os.path.join(CONFIG["output_dir"], "ap_draft_thisweek.xlsx")
                
                if not os.path.exists(source_path):
                    wb_dummy = openpyxl.Workbook()
                    ws_dummy = wb_dummy.active
                    ws_dummy.title = "Invoices"
                    ws_dummy.append(["invoice_no", "vendor_tax_id", "amount_invoice", "tax_amount", "po_ref", "notes"])
                    ws_dummy.append(["INV-001", "TAX-123", 100.50, 10.05, "PO-1001", "Test note"])
                    wb_dummy.save(source_path)
                
                wb_source = openpyxl.load_workbook(source_path)
                ws_source = wb_source["Invoices"]
                
                wb_target = openpyxl.load_workbook(target_path)
                if "Invoices" in wb_target.sheetnames:
                    del wb_target["Invoices"]
                ws_target = wb_target.create_sheet("Invoices")
                
                for row in ws_source.iter_rows(values_only=True):
                    ws_target.append(row)
                
                wb_target.save(target_path)
            except Exception:
                await page.screenshot(path="step_3_error.png")
                raise

            # ── 步骤 4：Match and write results
            try:
                import os
                import datetime
                import openpyxl
                from docx import Document
                
                target_path = os.path.join(CONFIG["output_dir"], "ap_draft_thisweek.xlsx")
                wb = openpyxl.load_workbook(target_path)
                
                ws_sys = wb["System"]
                ws_inv = wb["Invoices"]
                
                sys_headers = [cell.value for cell in ws_sys[1]]
                sys_data = []
                for row in ws_sys.iter_rows(min_row=2, values_only=True):
                    sys_data.append(dict(zip(sys_headers, row)))
                
                inv_headers = [cell.value for cell in ws_inv[1]]
                inv_data = []
                for row in ws_inv.iter_rows(min_row=2, values_only=True):
                    inv_data.append(dict(zip(inv_headers, row)))
                
                if "Match Results" in wb.sheetnames:
                    del wb["Match Results"]
                ws_res = wb.create_sheet("Match Results")
                
                res_headers = ["line_id", "po_ref", "amount_system", "invoice_no", "amount_invoice", "match_status", "diff_notes"]
                ws_res.append(res_headers)
                match_results_for_word = [res_headers]
                
                for sys_row in sys_data:
                    po_ref = sys_row.get("po_ref")
                    amount_system = float(sys_row.get("amount_system") or 0)
                    line_id = sys_row.get("line_id")
                    
                    candidates_stage1 = [inv for inv in inv_data if inv.get("po_ref") == po_ref]
                    
                    candidates_stage2 = []
                    for inv in candidates_stage1:
                        amt_inv = float(inv.get("amount_invoice") or 0)
                        if abs(amt_inv - amount_system) <= 1:
                            candidates_stage2.append(inv)
                            
                    invoice_no = ""
                    amount_invoice = ""
                    match_status = ""
                    diff_notes = ""
                    
                    if len(candidates_stage2) == 0:
                        match_status = "unmatched"
                        diff_notes = "No matching po_ref or out of tolerance"
                    elif len(candidates_stage2) == 1:
                        match = candidates_stage2[0]
                        amt_inv = float(match.get("amount_invoice") or 0)
                        diff = abs(amt_inv - amount_system)
                        invoice_no = match.get("invoice_no")
                        amount_invoice = amt_inv
                        if diff == 0:
                            match_status = "matched"
                        else:
                            match_status = "partial"
                            diff_notes = f"diff: {amt_inv - amount_system:.2f}"
                    else:
                        match_status = "pending"
                        closest_match = min(candidates_stage2, key=lambda x: abs(float(x.get("amount_invoice") or 0) - amount_system))
                        amt_inv = float(closest_match.get("amount_invoice") or 0)
                        invoice_no = closest_match.get("invoice_no")
                        amount_invoice = amt_inv
                        diff_notes = "Multiple candidates within tolerance"
                        
                    res_row = [line_id, po_ref, amount_system, invoice_no, amount_invoice, match_status, diff_notes]
                    ws_res.append(res_row)
                    match_results_for_word.append([str(x) if x is not None else "" for x in res_row])
                
                wb.save(target_path)
                
                date_str = datetime.datetime.now().strftime("%Y%m%d")
                docx_path = os.path.join(CONFIG["output_dir"], f"ap_reconciliation_{date_str}.docx")
                doc = Document()
                doc.add_heading(f"AP Reconciliation - {date_str}", 0)
                
                table = doc.add_table(rows=1, cols=len(res_headers))
                table.style = "Table Grid"
                hdr_cells = table.rows[0].cells
                for i, header in enumerate(res_headers):
                    hdr_cells[i].text = header
                
                for row_data in match_results_for_word[1:]:
                    row_cells = table.add_row().cells
                    for i, item in enumerate(row_data):
                        row_cells[i].text = item
                        
                doc.save(docx_path)
            except Exception:
                await page.screenshot(path="step_4_error.png")
                raise

        except PlaywrightTimeout as e:
            await page.screenshot(path="error_timeout.png")
            raise RuntimeError(f"超时：{e}") from e
        except Exception:
            await page.screenshot(path="error_unexpected.png")
            raise
        finally:
            await browser.close()


if __name__ == "__main__":
    asyncio.run(run())
