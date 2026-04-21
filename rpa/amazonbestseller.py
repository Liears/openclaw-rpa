#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Amazon Clothing, Shoes & Jewelry — 前50商品抓取
提取：标题、价格、评分、评论数、网址链接url
输出：追加到 ~/Desktop/amazon.docx（不存在则新建）

依赖：
  pip install playwright python-docx
  playwright install chromium
"""

import re
import time
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import docx.opc.constants
from playwright.sync_api import sync_playwright


# ─── 配置 ───────────────────────────────────────────────────────────────────
TARGET_URL  = "https://www.amazon.com/s?k=Clothing%2C+Shoes+%26+Jewelry&language=zh"
OUTPUT_PATH = Path("~/Desktop/amazon.docx").expanduser()
WANT_COUNT  = 50
HEADLESS    = False   # True = 无界面后台运行；False = 显示浏览器窗口


# ─── 工具函数 ─────────────────────────────────────────────────────────────────

def now_cn() -> str:
    t = datetime.now()
    return f"{t.month:02d}月{t.day:02d}日{t.hour:02d}时{t.minute:02d}分"


def clean(s: str) -> str:
    return re.sub(r"\s+", " ", s).strip() if s else ""


# ─── 页面提取 ─────────────────────────────────────────────────────────────────

def extract_page(page) -> list[dict]:
    """从当前页面提取所有搜索结果商品。"""
    try:
        page.wait_for_selector('[data-component-type="s-search-result"]', timeout=15_000)
    except Exception:
        return []

    items = page.query_selector_all('[data-component-type="s-search-result"]')
    products = []

    for item in items:
        try:
            # ── 标题：[data-cy="title-recipe"] ──
            title_el = item.query_selector('[data-cy="title-recipe"]')
            title = clean(title_el.inner_text()) if title_el else ""
            if not title:
                continue

            # ── URL：找第一个包含 /dp/ 的 a 标签 ──
            # Amazon 语言版 URL 格式: /-/zh/dp/ASIN/ref=... 或 /dp/ASIN/ref=...
            url = ""
            link_els = item.query_selector_all("a[href]")
            for link_el in link_els:
                href = link_el.get_attribute("href") or ""
                if "/dp/" in href:
                    full = ("https://www.amazon.com" + href) if href.startswith("/") else href
                    # 提取 ASIN，构造干净的标准商品链接
                    m = re.search(r"/dp/([A-Z0-9]{10})", full)
                    url = f"https://www.amazon.com/dp/{m.group(1)}" if m else full.split("?")[0]
                    break

            # ── 价格：[data-cy="price-recipe"] ──
            price = ""
            price_el = item.query_selector('[data-cy="price-recipe"] .a-offscreen')
            if price_el:
                price = clean(price_el.inner_text())
            else:
                price_block = item.query_selector('[data-cy="price-recipe"]')
                if price_block:
                    txt = clean(price_block.inner_text())
                    # 找第一个价格串：$XX.XX / JPY X,XXX / ¥XX 等
                    m = re.search(r"([$¥￥JPY]+\s*[\d,]+\.?\d*|[\d,]+\.?\d*\s*[$¥￥])", txt)
                    price = m.group(1) if m else txt[:20]

            # ── 评分：[data-cy="reviews-ratings-slot"] ──
            rating = ""
            rat_el = item.query_selector('[data-cy="reviews-ratings-slot"]')
            if rat_el:
                txt = clean(rat_el.inner_text())
                m = re.search(r"(\d+\.?\d*)", txt)
                rating = m.group(1) if m else ""

            # ── 评论数：reviews-block 里括号内的数字 ──
            reviews = ""
            rev_block = item.query_selector('[data-cy="reviews-block"]')
            if rev_block:
                txt = clean(rev_block.inner_text())
                m = re.search(r"\(([0-9,，]+)\)", txt)
                if m:
                    reviews = m.group(1).replace(",", "").replace("，", "")

            products.append({
                "title":   title,
                "price":   price   or "—",
                "rating":  rating  or "—",
                "reviews": reviews or "—",
                "url":     url     or "—",
            })
        except Exception:
            continue

    return products


# ─── Word 超链接辅助 ──────────────────────────────────────────────────────────

def add_hyperlink_to_cell(cell, url: str, text: str) -> None:
    """在 Word 表格单元格里插入可点击超链接。"""
    para = cell.paragraphs[0]
    # 向文档关系表注册超链接
    part = para.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run_elem = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")

    # 蓝色 + 下划线样式
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rpr.append(color)

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rpr.append(u)

    run_elem.append(rpr)

    t = OxmlElement("w:t")
    t.text = text
    run_elem.append(t)
    hyperlink.append(run_elem)
    para._p.append(hyperlink)


# ─── Word 写入 ────────────────────────────────────────────────────────────────

def append_to_word(products: list[dict], query_time: str, path: Path) -> None:
    """将结果追加到 Word 文档末尾（不存在则新建）。"""
    if path.exists():
        doc = Document(str(path))
        doc.add_page_break()
    else:
        doc = Document()

    # 查询时间段落
    p = doc.add_paragraph()
    run = p.add_run(f"查询时间：{query_time}")
    run.bold = True
    run.font.size = Pt(12)

    # 表格
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"

    # 表头
    headers = ["标题", "价格", "评分", "评论数", "网址链接url"]
    hdr_row = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_row[i].text = h
        run = hdr_row[i].paragraphs[0].runs[0]
        run.bold = True

    # 数据行
    for prod in products:
        row = table.add_row().cells
        row[0].text = prod["title"]
        row[1].text = prod["price"]
        row[2].text = prod["rating"]
        row[3].text = prod["reviews"]
        # URL 列：可点击超链接（如果是 — 则纯文本）
        url = prod["url"]
        if url and url != "—":
            add_hyperlink_to_cell(row[4], url, url)
        else:
            row[4].text = url

    # 列宽：标题宽一些，URL 宽一些
    col_widths = [Inches(2.5), Inches(0.8), Inches(0.5), Inches(0.8), Inches(2.0)]
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = col_widths[i]

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    print(f"✅ 已保存 {len(products)} 条记录 → {path}")


# ─── 主流程 ───────────────────────────────────────────────────────────────────

def main():
    query_time = now_cn()
    print(f"⏱  查询时间: {query_time}")
    print(f"🌐 目标: {TARGET_URL}")
    print(f"📄 输出: {OUTPUT_PATH}\n")

    all_products: list[dict] = []

    with sync_playwright() as pw:
        browser = pw.chromium.launch(
            headless=HEADLESS,
            args=["--disable-blink-features=AutomationControlled"],
        )
        ctx = browser.new_context(
            viewport={"width": 1440, "height": 900},
            user_agent=(
                "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
                "AppleWebKit/537.36 (KHTML, like Gecko) "
                "Chrome/124.0.0.0 Safari/537.36"
            ),
            locale="zh-CN",
            extra_http_headers={"Accept-Language": "zh-CN,zh;q=0.9,en;q=0.8"},
        )
        # 隐藏 webdriver 标识，降低被拦截概率
        ctx.add_init_script(
            "Object.defineProperty(navigator,'webdriver',{get:()=>undefined})"
        )
        page = ctx.new_page()

        page_num = 1
        current_url = TARGET_URL

        while len(all_products) < WANT_COUNT:
            print(f"  📃 第 {page_num} 页 → {current_url[:80]}...")
            try:
                page.goto(current_url, wait_until="domcontentloaded", timeout=30_000)
                time.sleep(2)  # 让 JS 稳定渲染
            except Exception as e:
                print(f"  ⚠️  页面加载失败: {e}")
                break

            # 检测是否被验证码拦截
            if page.query_selector("form[action*='captcha']") or "captcha" in page.url.lower():
                print("  ⚠️  遇到验证码，请在浏览器中手动通过后按 Enter 继续...")
                input()

            items = extract_page(page)
            if not items:
                print("  ⚠️  本页未提取到商品，停止翻页。")
                break

            before = len(all_products)
            all_products.extend(items)
            print(f"  ✓  本页 {len(items)} 条，累计 {len(all_products)} 条")

            if len(all_products) >= WANT_COUNT:
                break

            # 找下一页按钮
            next_el = page.query_selector("a.s-pagination-next:not(.s-pagination-disabled)")
            if not next_el:
                print("  ℹ️  已无下一页。")
                break

            next_href = next_el.get_attribute("href") or ""
            current_url = ("https://www.amazon.com" + next_href) if next_href.startswith("/") else next_href
            page_num += 1
            time.sleep(1.5)

        browser.close()

    # 截取前 50 条
    all_products = all_products[:WANT_COUNT]

    if not all_products:
        print("❌ 未抓取到任何数据，请检查网络或 Amazon 页面结构。")
        return

    print(f"\n共提取 {len(all_products)} 条商品，写入 Word...")
    append_to_word(all_products, query_time, OUTPUT_PATH)


if __name__ == "__main__":
    main()
