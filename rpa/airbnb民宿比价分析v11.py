# pip install playwright httpx openpyxl python-docx && playwright install chromium
# 任务：Airbnb民宿比价分析V11
# 录制时间：2026-04-13 16:29:05
# 由 OpenClaw RPA Recorder（headed 真实录制）生成 — 可脱离 OpenClaw 独立运行

import asyncio
import datetime
import json
import os
import re
import urllib.parse
from pathlib import Path

import httpx
from docx import Document

from playwright.async_api import async_playwright, TimeoutError as PlaywrightTimeout

CONFIG = {
    "output_dir":    Path.home() / "Desktop",
    "headless":      False,
    "timeout":       60_000,
    "slow_mo":       300,
    # 导航后等待 SPA 内容渲染的额外时间（重型 SPA 如 Yahoo Finance 需要 1-2 秒）
    "spa_settle_ms": 1_500,
    # extract_text 等待目标元素出现的超时（毫秒）
    "content_wait":  15_000,
    # httpx 调用外部 API 的超时（秒）
    "api_timeout":   60.0,
    # extract_text / snapshot / dom_inspect 前等待主内容（骨架屏后再读 DOM）
    "extract_ready_timeout_ms": 30_000,
    # 若有已保存的登录 Cookie，填写路径（由 rpa_manager login-start 生成）
    "cookies_path":  "",    # ── 视觉识别配置（由 OpenClaw RPA 录制时自动生成）──────────────────────────
    "vision_model":    'qwen3-vl-plus',
    "vision_base_url": 'https://dashscope.aliyuncs.com/compatible-mode/v1',
    "vision_api_key":  'sk-e947b516a5b24e889976ef78314646e6',   # DASHSCOPE_API_KEY
    "vision_ready_timeout_ms": 45_000,  # 视觉截图前等待；Airbnb 等可调 60_000

}

_UA = (
    "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
    "AppleWebKit/537.36 (KHTML, like Gecko) "
    "Chrome/124.0.0.0 Safari/537.36"
)

def _write_kv_field(out_path: Path, field_name: str, values: list[str], first_write: bool) -> None:
    """将 extract_text 提取结果写入 kv 格式临时文件。

    格式规则（兼容中英文字段名与字段值）：
      单值：  field_name: value
      多值：  field_name.0: value0
              field_name.1: value1
              ...
    追加写入时直接续行；调用方负责 first_write 标志。
    """
    if not values:
        return  # 0 条时不写任何内容（read 时 raise RuntimeError 提示选择器问题）
    lines: list[str] = []
    if len(values) == 1:
        lines.append(f"{field_name}: {values[0]}")
    else:
        for i, v in enumerate(values):
            lines.append(f"{field_name}.{i}: {v}")
    blob = "\n".join(lines) + "\n"
    if first_write:
        out_path.write_text(blob, encoding="utf-8")
    else:
        with out_path.open("a", encoding="utf-8") as f:
            f.write(blob)


def _parse_field(filepath, field_name: str, index: int = 0):
    """从 extract_text 输出的 kv 文件中读取指定字段值。

    兼容中英文字段名与值（UTF-8）。支持带 .N 索引后缀的多值字段。

    Args:
        filepath:   文件路径（str 或 Path）
        field_name: 字段名，与 extract_text 的 field 参数一致
        index:      0 = 第一条（默认）；-1 = 最后一条；None = 返回全部列表

    Raises:
        RuntimeError: 文件不存在，或字段在文件中未找到
    """
    path = Path(filepath) if not isinstance(filepath, Path) else filepath
    if not path.exists():
        raise RuntimeError(
            f"提取文件不存在 / Extract file not found: {path}\n"
            f"请确认 extract_text 步骤已成功执行并写入该文件 / "
            f"Make sure the extract_text step ran successfully and wrote this file."
        )
    matches: list[str] = []
    for line in path.read_text(encoding="utf-8").splitlines():
        line = line.strip()
        if not line or line.startswith("#"):
            continue
        if ":" not in line:
            continue
        k, _, v = line.partition(":")
        k_base = k.strip()
        # 去掉 .N 索引后缀：field.0 → field
        if "." in k_base and k_base.rsplit(".", 1)[-1].isdigit():
            k_base = k_base.rsplit(".", 1)[0]
        if k_base == field_name:
            matches.append(v.strip())
    if not matches:
        raise RuntimeError(
            f"字段 '{field_name}' 在 {path} 中未找到 / "
            f"Field '{field_name}' not found in {path}.\n"
            f"请检查 extract_text 的 field 参数是否与此处一致 / "
            f"Check that the extract_text 'field' param matches this name."
        )
    if index is None:
        return matches
    try:
        return matches[index]
    except IndexError:
        return matches[-1]


_EXTRACT_JS = '([s,n])=>{const r=document.querySelector("main")||document.querySelector(\'[role="main"]\');const bare=/^[a-zA-Z][a-zA-Z0-9-]*$/.test(s)&&s.indexOf("#")<0&&s.indexOf(".")<0&&s.indexOf("[")<0&&s.indexOf(" ")<0;const sc=bare&&r?r:document;return Array.from(sc.querySelectorAll(s)).slice(0,n).map(e=>(e.textContent||"").replace(/\\s+/g," ").trim()).filter(Boolean)}'
async def _wait_spa_ready_for_vision(
    page,
    crop_selector: str = "",
    *,
    timeout_ms: int = 45_000,
) -> None:
    """视觉截图前等待 SPA 主内容就绪，减少骨架屏、未 hydration 就截图的情况。

    顺序：domcontentloaded → 尝试 networkidle → 固定短等 → 轮询「大图已解码或正文足够」；
    轮询中偶尔 wheel 触发懒加载。若提供 crop_selector，最后再等该容器 visible。
    """
    import time as _time

    try:
        await page.wait_for_load_state("domcontentloaded", timeout=12_000)
    except Exception:
        pass
    try:
        await page.wait_for_load_state("networkidle", timeout=28_000)
    except Exception:
        pass
    await page.wait_for_timeout(1_200)

    deadline = _time.monotonic() + max(5_000, timeout_ms) / 1000.0
    poll_js = """() => {
        const imgs = Array.from(document.querySelectorAll('img'));
        for (const i of imgs) {
            if (i.complete && i.naturalWidth > 64 && i.naturalHeight > 64) return true;
        }
        const t = (document.body && document.body.innerText) || '';
        const compact = t.replace(/\\s+/g, '');
        if (compact.length > 380 && /\\d/.test(t)) {
            if (/[\\u4e00-\\u9fff]/.test(t) || compact.length > 620) return true;
        }
        return false;
    }"""
    n = 0
    while _time.monotonic() < deadline:
        try:
            if await page.evaluate(poll_js):
                await page.wait_for_timeout(700)
                break
        except Exception:
            pass
        n += 1
        if n % 5 == 0:
            try:
                await page.mouse.wheel(0, 320)
            except Exception:
                pass
        await page.wait_for_timeout(420)
    else:
        await page.wait_for_timeout(1_000)

    if crop_selector and str(crop_selector).strip():
        try:
            await page.locator(crop_selector.strip()).first.wait_for(
                state="visible", timeout=min(25_000, timeout_ms)
            )
            await page.wait_for_timeout(400)
        except Exception:
            pass



async def _vision_call(fields: list, image_bytes: bytes, cfg: dict) -> dict:
    """调用视觉 LLM 从截图提取字段。模型/Key 由 CONFIG 注入，无需手动配置。"""
    import base64 as _vb64, re as _vre
    b64 = _vb64.b64encode(image_bytes).decode()
    tmpl = {f: "" for f in fields}
    prompt = (
        f"从截图中提取以下字段，只返回 JSON，不要解释：\n"
        + __import__("json").dumps(tmpl, ensure_ascii=False)
        + "\n规则：①只提取可见文字；②看不到设为空字符串；③价格保留原始格式。"
    )
    payload = {
        "model": cfg["vision_model"],
        "messages": [{"role": "user", "content": [
            {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{b64}"}},
            {"type": "text", "text": prompt},
        ]}],
        "max_tokens": 500,
    }
    base = cfg["vision_base_url"].rstrip("/")
    _vt = httpx.Timeout(300.0, connect=30.0)
    async with httpx.AsyncClient(timeout=_vt, verify=False) as _hc:
        _r = await _hc.post(
            f"{base}/chat/completions",
            headers={"Authorization": f"Bearer {cfg['vision_api_key']}", "Content-Type": "application/json"},
            json=payload,
        )
        _r.raise_for_status()
    raw = _r.json()["choices"][0]["message"]["content"].strip()
    raw = _vre.sub(r"^```(?:json)?\s*|\s*```$", "", raw, flags=_vre.MULTILINE).strip()
    return __import__("json").loads(raw)



async def _wait_for_content(page, selector: str) -> None:
    """等待 selector 对应的元素出现在 DOM 中（容错：超时也继续）。"""
    try:
        await page.wait_for_selector(selector, timeout=CONFIG["content_wait"])
    except Exception:
        pass  # 元素未出现也继续，evaluate 会返回空列表


async def _scroll_window(page, dy: int) -> None:
    """窗口滚动：导航后若再用 evaluate(scrollBy)，易因执行上下文销毁报错；用 mouse.wheel 并在滚动前等待页面稳定。"""
    try:
        await page.wait_for_load_state("domcontentloaded", timeout=10_000)
    except Exception:
        pass
    vp = page.viewport_size
    if vp:
        await page.mouse.move(vp["width"] // 2, vp["height"] // 2)
    else:
        await page.mouse.move(720, 450)
    await page.mouse.wheel(0, float(dy))


async def run():
    async with async_playwright() as p:
        browser = await p.chromium.launch(
            headless=CONFIG["headless"],
            args=["--no-sandbox", "--disable-blink-features=AutomationControlled"],
            slow_mo=CONFIG["slow_mo"],
        )
        context = await browser.new_context(
            user_agent=_UA,
            viewport={"width": 1440, "height": 900},
            locale="en-US",
            timezone_id="America/New_York",
            extra_http_headers={"Accept-Language": "en-US,en;q=0.9"},
        )
        await context.add_init_script(
            "Object.defineProperty(navigator, 'webdriver', {get: () => undefined})"
        )
        # 若配置了 cookies_path，注入 Cookie 模拟已登录态（由 rpa_manager login-start 生成）
        import json as _json
        _cp = CONFIG.get("cookies_path", "")
        if _cp and Path(_cp).exists():
            await context.add_cookies(_json.loads(Path(_cp).read_text()))
        page = await context.new_page()
        page.set_default_timeout(CONFIG["timeout"])

        try:
            # ── 步骤 1：步骤 1
            try:
                await page.goto('https://www.airbnb.cn/rooms/1517880824760006835?location=%C5%8Csaka-shi%2C%20%C5%8Csaka-fu%2C%20JP&search_mode=regular_search&adults=1&check_in=2026-04-20&check_out=2026-04-27', wait_until="domcontentloaded")
                await page.wait_for_timeout(CONFIG["spa_settle_ms"])
            except Exception:
                await page.screenshot(path="step_1_error.png")
                raise

            # ── 步骤 2：步骤 2
            try:
                # 视觉截图前等待页面就绪（骨架屏消失 / 大图或正文出现）
                await _wait_spa_ready_for_vision(page, '', timeout_ms=CONFIG["vision_ready_timeout_ms"])
                # 截图
                _vision_shot = CONFIG["output_dir"] / "vision_step_02.png"
                await page.screenshot(path=str(_vision_shot), full_page=False)
                _extracted = await _vision_call(['民宿名字', '房间名称', '价格', '评分'], _vision_shot.read_bytes(), CONFIG)
                _out = CONFIG["output_dir"] / '/tmp/rpa_hotel1.txt'
                _first_v = True
                for _i, _f in enumerate(['民宿名字', '房间名称', '价格', '评分']):
                    _val = str(_extracted.get(_f, "")).strip()
                    _write_kv_field(_out, _f, [_val] if _val else [], first_write=(_i == 0 and _first_v))
                print(f"视觉提取完成 → {_out}  字段：{list(_extracted.keys())}")
            except Exception:
                await page.screenshot(path="step_2_error.png")
                raise

            # ── 步骤 3：步骤 3
            try:
                await page.goto('https://www.airbnb.cn/rooms/1520897875971878894?location=%C5%8Csaka-shi%2C%20%C5%8Csaka-fu%2C%20JP&search_mode=regular_search&adults=1&check_in=2026-04-20&check_out=2026-04-27', wait_until="domcontentloaded")
                await page.wait_for_timeout(CONFIG["spa_settle_ms"])
            except Exception:
                await page.screenshot(path="step_3_error.png")
                raise

            # ── 步骤 4：步骤 4
            try:
                await page.wait_for_timeout(5000)
            except Exception:
                await page.screenshot(path="step_4_error.png")
                raise

            # ── 步骤 5：步骤 5
            try:
                # 视觉截图前等待页面就绪（骨架屏消失 / 大图或正文出现）
                await _wait_spa_ready_for_vision(page, '', timeout_ms=CONFIG["vision_ready_timeout_ms"])
                # 截图
                _vision_shot = CONFIG["output_dir"] / "vision_step_05.png"
                await page.screenshot(path=str(_vision_shot), full_page=False)
                _extracted = await _vision_call(['民宿名字', '房间名称', '价格', '评分'], _vision_shot.read_bytes(), CONFIG)
                _out = CONFIG["output_dir"] / '/tmp/rpa_hotel2.txt'
                _first_v = True
                for _i, _f in enumerate(['民宿名字', '房间名称', '价格', '评分']):
                    _val = str(_extracted.get(_f, "")).strip()
                    _write_kv_field(_out, _f, [_val] if _val else [], first_write=(_i == 0 and _first_v))
                print(f"视觉提取完成 → {_out}  字段：{list(_extracted.keys())}")
            except Exception:
                await page.screenshot(path="step_5_error.png")
                raise

            # ── 步骤 6：步骤 6
            try:
                await page.goto('https://www.airbnb.cn/rooms/1239558906468787551?check_in=2026-04-20&check_out=2026-04-27&location=%E6%96%B0%E4%BB%8A%E5%AE%AB%E7%AB%99&search_mode=regular_search&adults=1', wait_until="domcontentloaded")
                await page.wait_for_timeout(CONFIG["spa_settle_ms"])
            except Exception:
                await page.screenshot(path="step_6_error.png")
                raise

            # ── 步骤 7：步骤 7
            try:
                # 视觉截图前等待页面就绪（骨架屏消失 / 大图或正文出现）
                await _wait_spa_ready_for_vision(page, '', timeout_ms=CONFIG["vision_ready_timeout_ms"])
                # 截图
                _vision_shot = CONFIG["output_dir"] / "vision_step_07.png"
                await page.screenshot(path=str(_vision_shot), full_page=False)
                _extracted = await _vision_call(['民宿名字', '房间名称', '价格', '评分'], _vision_shot.read_bytes(), CONFIG)
                _out = CONFIG["output_dir"] / '/tmp/rpa_hotel3.txt'
                _first_v = True
                for _i, _f in enumerate(['民宿名字', '房间名称', '价格', '评分']):
                    _val = str(_extracted.get(_f, "")).strip()
                    _write_kv_field(_out, _f, [_val] if _val else [], first_write=(_i == 0 and _first_v))
                print(f"视觉提取完成 → {_out}  字段：{list(_extracted.keys())}")
            except Exception:
                await page.screenshot(path="step_7_error.png")
                raise

            # ── 步骤 8：步骤 8
            try:
                import json, datetime
                
                def _parse_field(file_path, field, index=0):
                    lines = Path(file_path).read_text(encoding='utf-8').strip().split('\n')
                    values = [l.split(':', 1)[1].strip() for l in lines if l.startswith(field + ':')]
                    if index is None:
                        return values
                    if index == -1:
                        return values[-1] if values else ''
                    return values[index] if index < len(values) else ''
                
                query_time = datetime.datetime.now().strftime('%m月%d日%H时%M分')
                ruzhu_time = '2026-04-20 到 2026-04-27'
                
                files = ['/tmp/rpa_hotel1.txt', '/tmp/rpa_hotel2.txt', '/tmp/rpa_hotel3.txt']
                headers = ['民宿名字', '房间名称', '价格', '评分', '查询时间', '入住时间']
                rows = []
                
                for f in files:
                    row = [
                        _parse_field(f, '民宿名字'),
                        _parse_field(f, '房间名称'),
                        _parse_field(f, '价格'),
                        _parse_field(f, '评分'),
                        query_time,
                        ruzhu_time
                    ]
                    rows.append(row)
                
                output = {'headers': headers, 'rows': rows}
                Path('/tmp/rpa_aggregate.json').write_text(json.dumps(output, ensure_ascii=False, indent=2), encoding='utf-8')
                print(f'Aggregated {len(rows)} rows to /tmp/rpa_aggregate.json')
            except Exception:
                await page.screenshot(path="step_8_error.png")
                raise

            # ── 步骤 9：步骤 9
            try:
                _wp = Path('~/Desktop/Airbnb/hotelCompare.docx').expanduser()
                _wp.parent.mkdir(parents=True, exist_ok=True)
                _wparas = []
                _wmode = 'append'
                if _wmode == "append" and _wp.exists():
                    _doc = Document(str(_wp))
                else:
                    _doc = Document()
                for _p in _wparas:
                    _doc.add_paragraph(str(_p))
                _doc.save(str(_wp))
                print("word_write →", _wp)
            except Exception:
                await page.screenshot(path="step_9_error.png")
                raise

        except PlaywrightTimeout as e:
            await page.screenshot(path="error_timeout.png")
            raise RuntimeError(f"超时：{e}") from e
        except Exception:
            await page.screenshot(path="error_unexpected.png")
            raise
        finally:
            await browser.close()


if __name__ == "__main__":
    asyncio.run(run())
