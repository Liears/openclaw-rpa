# pip install playwright httpx openpyxl python-docx && playwright install chromium
# 任务：会计记账V2
# 录制时间：2026-04-07 01:18:54
# 由 OpenClaw RPA Recorder（headed 真实录制）生成 — 可脱离 OpenClaw 独立运行

import asyncio
import os
import urllib.parse
from pathlib import Path

import httpx
from openpyxl import Workbook, load_workbook
from openpyxl.utils import get_column_letter

from docx import Document

from playwright.async_api import async_playwright, TimeoutError as PlaywrightTimeout

CONFIG = {
    "output_dir":    Path.home() / "Desktop",
    "headless":      False,
    "timeout":       60_000,
    "slow_mo":       300,
    # 导航后等待 SPA 内容渲染的额外时间（重型 SPA 如 Yahoo Finance 需要 1-2 秒）
    "spa_settle_ms": 1_500,
    # extract_text 等待目标元素出现的超时（毫秒）
    "content_wait":  15_000,
    # httpx 调用外部 API 的超时（秒）
    "api_timeout":   60.0,}

_UA = (
    "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
    "AppleWebKit/537.36 (KHTML, like Gecko) "
    "Chrome/124.0.0.0 Safari/537.36"
)

def _format_extract_section(field_label: str, lines: list[str]) -> str:
    """Format extracted DOM text: show field name, separator line, then body."""
    name = (field_label or "").strip() or "extract"
    title = f"【字段：{name}】"
    if not lines:
        body = "(no text matched)\n"
    elif len(lines) == 1:
        body = lines[0].strip() + "\n"
    else:
        parts = [f"{i + 1}. {s.strip()}" for i, s in enumerate(lines)]
        body = "\n\n".join(parts) + "\n"
    sep = "─" * 32
    return f"{title}\n{sep}\n{body}\n"


_EXTRACT_JS = '([s,n])=>{const r=document.querySelector("main")||document.querySelector(\'[role="main"]\');const bare=/^[a-zA-Z][a-zA-Z0-9-]*$/.test(s)&&s.indexOf("#")<0&&s.indexOf(".")<0&&s.indexOf("[")<0&&s.indexOf(" ")<0;const sc=bare&&r?r:document;return Array.from(sc.querySelectorAll(s)).slice(0,n).map(e=>(e.textContent||"").replace(/\\s+/g," ").trim()).filter(Boolean)}'


async def _wait_for_content(page, selector: str) -> None:
    """等待 selector 对应的元素出现在 DOM 中（容错：超时也继续）。"""
    try:
        await page.wait_for_selector(selector, timeout=CONFIG["content_wait"])
    except Exception:
        pass  # 元素未出现也继续，evaluate 会返回空列表


async def _scroll_window(page, dy: int) -> None:
    """窗口滚动：导航后若再用 evaluate(scrollBy)，易因执行上下文销毁报错；用 mouse.wheel 并在滚动前等待页面稳定。"""
    try:
        await page.wait_for_load_state("domcontentloaded", timeout=10_000)
    except Exception:
        pass
    vp = page.viewport_size
    if vp:
        await page.mouse.move(vp["width"] // 2, vp["height"] // 2)
    else:
        await page.mouse.move(720, 450)
    await page.mouse.wheel(0, float(dy))


async def run():
    async with async_playwright() as p:
        browser = await p.chromium.launch(
            headless=CONFIG["headless"],
            args=["--no-sandbox", "--disable-blink-features=AutomationControlled"],
            slow_mo=CONFIG["slow_mo"],
        )
        context = await browser.new_context(
            user_agent=_UA,
            viewport={"width": 1440, "height": 900},
            locale="en-US",
            timezone_id="America/New_York",
            extra_http_headers={"Accept-Language": "en-US,en;q=0.9"},
        )
        await context.add_init_script(
            "Object.defineProperty(navigator, 'webdriver', {get: () => undefined})"
        )
        page = await context.new_page()
        page.set_default_timeout(CONFIG["timeout"])

        try:
            # ── 步骤 1：拉取本周应付待对账数据（跳过SSL校验）
            try:
                _params = {'status': 'open', 'week': '2026-W14'}
                _api_url = 'https://0a34723da37946b7add0b4581c37ada2_oas.api.mockbin.io/ap/reconciliation/batches' + "?" + urllib.parse.urlencode(_params)
                async with httpx.AsyncClient(timeout=CONFIG["api_timeout"], verify=False) as _hc:
                    _r = await _hc.get(_api_url)
                    _r.raise_for_status()
                (CONFIG["output_dir"] / 'reconcile_raw.json').write_text(_r.text, encoding="utf-8")
                print("API 响应已写入", CONFIG["output_dir"] / 'reconcile_raw.json')
            except Exception:
                await page.screenshot(path="step_1_error.png")
                raise

            # ── 步骤 2：读取 reconcile_raw.json 展平后写入「系统侧」工作表
            try:
                from pathlib import Path
                from openpyxl import Workbook, load_workbook
                import json
                
                output_dir = Path.home() / "Desktop"
                fpath = output_dir / "对账底稿_本周.xlsx"
                raw = json.loads((output_dir / "reconcile_raw.json").read_text())
                
                rows = []
                for batch in raw.get("batches", []):
                    for line in batch.get("lines", []):
                        rows.append({
                            "line_id": line["line_id"],
                            "vendor_id": line["vendor_id"],
                            "po_ref": line["po_ref"],
                            "amount_system": line["amount_system"],
                            "currency": line["currency"],
                            "due_date": line["due_date"],
                            "batch_id": batch["batch_id"],
                        })
                
                headers = ["line_id", "vendor_id", "po_ref", "amount_system", "currency", "due_date", "batch_id"]
                
                if fpath.exists():
                    wb = load_workbook(fpath)
                else:
                    wb = Workbook()
                    wb.remove(wb.active)
                
                if "系统侧" in wb.sheetnames:
                    del wb["系统侧"]
                ws = wb.create_sheet("系统侧")
                
                for c, h in enumerate(headers, 1):
                    ws.cell(row=1, column=c, value=h)
                
                for r_idx, row in enumerate(rows, 2):
                    for c, key in enumerate(headers, 1):
                        ws.cell(row=r_idx, column=c, value=row[key])
                
                ws.freeze_panes = "A2"
                wb.save(fpath)
                print("已保存:", fpath, "工作表[系统侧]", len(rows), "条数据，冻结首行")
            except Exception:
                await page.screenshot(path="step_2_error.png")
                raise

            # ── 步骤 3：读取发票数据写入「发票侧」工作表
            try:
                from pathlib import Path
                from openpyxl import load_workbook
                
                output_dir = Path.home() / "Desktop"
                src_path = output_dir / "发票导入_本周.xlsx"
                dst_path = output_dir / "对账底稿_本周.xlsx"
                
                wb_src = load_workbook(src_path)
                ws_src = wb_src["发票侧"]
                rows = list(ws_src.iter_rows(values_only=True))
                wb_src.close()
                
                wb_dst = load_workbook(dst_path)
                if "发票侧" in wb_dst.sheetnames:
                    del wb_dst["发票侧"]
                ws_dst = wb_dst.create_sheet("发票侧")
                
                for row in rows:
                    ws_dst.append(list(row))
                
                wb_dst.save(dst_path)
                print("已写入:", dst_path, "[发票侧]", len(rows)-1, "条")
            except Exception:
                await page.screenshot(path="step_3_error.png")
                raise

            # ── 步骤 4：本地匹配（po_ref容差1元）+ 匹配结果sheet + Word报告
            try:
                from pathlib import Path
                from openpyxl import load_workbook
                from openpyxl.styles import PatternFill
                from docx import Document
                from datetime import date
                
                output_dir = Path.home() / "Desktop"
                dst_path = output_dir / "对账底稿_本周.xlsx"
                
                wb = load_workbook(dst_path)
                
                # 1. Read system side
                ws_sys = wb["系统侧"]
                system_rows = []
                for i, row in enumerate(ws_sys.iter_rows(values_only=True)):
                    if i == 0: continue
                    system_rows.append({"line_id": row[0], "po_ref": row[2], "amount": float(row[3])})
                
                # 2. Read invoice side
                ws_inv = wb["发票侧"]
                invoice_rows = []
                for i, row in enumerate(ws_inv.iter_rows(values_only=True)):
                    if i == 0: continue
                    invoice_rows.append({"invoice_no": row[0], "amount": float(row[2]), "po_ref": row[4]})
                
                print("系统侧:", len(system_rows), "行  发票侧:", len(invoice_rows), "行")
                
                # 3. Match with tolerance check
                matched_inv_indices = set()
                results = []
                
                for si, sys in enumerate(system_rows):
                    raw_candidates = [(ii, inv, abs(inv["amount"] - sys["amount"]))
                                  for ii, inv in enumerate(invoice_rows)
                                  if inv["po_ref"] == sys["po_ref"] and ii not in matched_inv_indices]
                    # Apply tolerance filter
                    candidates = [(ii, inv, d) for ii, inv, d in raw_candidates if d <= 1]
                
                    if not candidates:
                        if raw_candidates:
                            diffs = ", ".join(str(round(d,2)) for _, _, d in raw_candidates)
                            results.append({"line_id": sys["line_id"], "po_ref": sys["po_ref"],
                                            "sys_amount": sys["amount"], "invoice_no": "", "inv_amount": "",
                                            "status": "unmatched",
                                            "diff": "PO匹配但金额差异均超容差: " + diffs})
                        else:
                            results.append({"line_id": sys["line_id"], "po_ref": sys["po_ref"],
                                            "sys_amount": sys["amount"], "invoice_no": "", "inv_amount": "",
                                            "status": "unmatched", "diff": "系统侧无匹配发票"})
                    elif len(candidates) == 1:
                        ii, inv, diff = candidates[0]
                        matched_inv_indices.add(ii)
                        if diff == 0:
                            status = "matched"
                            diff_note = ""
                        else:
                            status = "partial"
                            diff_note = "差额 " + str(round(diff, 2)) + " 元"
                        results.append({"line_id": sys["line_id"], "po_ref": sys["po_ref"],
                                        "sys_amount": sys["amount"], "invoice_no": inv["invoice_no"],
                                        "inv_amount": inv["amount"], "status": status, "diff": diff_note})
                    else:
                        results.append({"line_id": sys["line_id"], "po_ref": sys["po_ref"],
                                        "sys_amount": sys["amount"], "invoice_no": "多候选", "inv_amount": "",
                                        "status": "pending",
                                        "diff": str(len(candidates)) + " 张发票在容差内匹配同一PO"})
                
                for ii, inv in enumerate(invoice_rows):
                    if ii not in matched_inv_indices:
                        results.append({"line_id": "", "po_ref": inv["po_ref"] or "（无PO）",
                                        "sys_amount": "", "invoice_no": inv["invoice_no"],
                                        "inv_amount": inv["amount"], "status": "unmatched",
                                        "diff": "系统侧无对应行"})
                
                print("匹配结果:", len(results), "行")
                for r in results: print(r)
                
                # 4. Write 匹配结果 sheet
                if "匹配结果" in wb.sheetnames:
                    del wb["匹配结果"]
                ws_res = wb.create_sheet("匹配结果")
                headers = ["系统行ID", "订单引用", "系统金额", "发票号码", "发票金额", "匹配状态", "差异说明"]
                ws_res.append(headers)
                status_colors = {"matched": "C6EFCE", "partial": "FFEB9C", "unmatched": "FFC7CE", "pending": "DDEBF7"}
                for r_idx, row in enumerate(results, 2):
                    ws_res.append([row["line_id"], row["po_ref"], row["sys_amount"],
                                   row["invoice_no"], row["inv_amount"], row["status"], row["diff"]])
                    color = status_colors.get(row["status"], "FFFFFF")
                    for c in range(1, 8):
                        ws_res.cell(row=r_idx, column=c).fill = PatternFill("solid", fgColor=color)
                wb.save(dst_path)
                print("匹配结果 sheet 写入完成")
                
                # 5. Word report
                report_date = date.today().strftime("%Y%m%d")
                doc = Document()
                doc.add_heading("应付对账报告", 0)
                doc.add_paragraph("生成日期：" + date.today().strftime("%Y-%m-%d") + "  批次：batch_2026w14_01  期间：2026-W14")
                doc.add_heading("匹配结果汇总", 1)
                table = doc.add_table(rows=1, cols=7)
                table.style = "Table Grid"
                hdr = table.rows[0].cells
                for i, h in enumerate(headers):
                    hdr[i].text = h
                    hdr[i].paragraphs[0].runs[0].font.bold = True
                for row in results:
                    cells = table.add_row().cells
                    for i, key in enumerate(["line_id","po_ref","sys_amount","invoice_no","inv_amount","status","diff"]):
                        cells[i].text = str(row[key])
                doc.add_paragraph("")
                matched_cnt = sum(1 for r in results if r["status"]=="matched")
                partial_cnt = sum(1 for r in results if r["status"]=="partial")
                unmatched_cnt = sum(1 for r in results if r["status"]=="unmatched")
                pending_cnt = sum(1 for r in results if r["status"]=="pending")
                summary = "共 " + str(len(results)) + " 条记录，matched " + str(matched_cnt) + " 条，partial " + str(partial_cnt) + " 条，unmatched " + str(unmatched_cnt) + " 条，pending " + str(pending_cnt) + " 条。"
                doc.add_paragraph(summary)
                report_path = output_dir / ("对账报告_" + report_date + ".docx")
                doc.save(report_path)
                print("Word 报告已保存:", report_path)
            except Exception:
                await page.screenshot(path="step_4_error.png")
                raise

        except PlaywrightTimeout as e:
            await page.screenshot(path="error_timeout.png")
            raise RuntimeError(f"超时：{e}") from e
        except Exception:
            await page.screenshot(path="error_unexpected.png")
            raise
        finally:
            await browser.close()


if __name__ == "__main__":
    asyncio.run(run())
