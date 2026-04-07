#!/usr/bin/env python3
"""
RPA Recorder Server
===================
Long-running headed Playwright process started by `rpa_manager record-start`.

IPC protocol (file-based):
  rpa_manager writes  → SESSION_DIR/cmd.json   {"action":..., "seq": N, ...}
  recorder_server writes → SESSION_DIR/result_N.json  {"success":..., "snapshot":[], ...}

Each executed action:
  - Runs in the headed (visible) Chromium window (browser steps), or runs locally for
    `api_call`, `merge_files`, `excel_write`, `word_write` (still snapshots the page after)
  - Takes a screenshot
  - Returns DOM snapshot so the LLM can pick real CSS selectors
  - Appends generated Python code to code_blocks list

On shutdown:
  - Compiles code_blocks into a full standalone Playwright script
  - Saves to SESSION_DIR/script_log.py
  - Writes SESSION_DIR/done marker
"""

import asyncio
import inspect
import json
import os
import re
import sys
import urllib.parse
from datetime import datetime
from pathlib import Path

import httpx

# ── 时间戳日志：所有 print() 自动加前缀 [HH:MM:SS] ──────────────────────────
# Timestamped logging: prepend [HH:MM:SS] to every print() call
import builtins as _builtins
_real_print = _builtins.print
def _timed_print(*args, **kwargs):  # type: ignore[override]
    _real_print(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}]", *args, **kwargs)
_builtins.print = _timed_print
# ─────────────────────────────────────────────────────────────────────────────

SKILL_DIR   = Path(__file__).parent
SESSION_DIR = SKILL_DIR / "recorder_session"

POLL_INTERVAL = 0.15  # seconds

# 同一次录制会话内，同一输出文件名多次 extract_text：首次 write_text，之后 open("a") 追加，避免生成脚本互相覆盖
_EXTRACT_OUTPUT_FILES: set[str] = set()


def _reset_extract_output_tracking() -> None:
    global _EXTRACT_OUTPUT_FILES
    _EXTRACT_OUTPUT_FILES = set()


_UA = (
    "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
    "AppleWebKit/537.36 (KHTML, like Gecko) "
    "Chrome/124.0.0.0 Safari/537.36"
)

# 与 _build_final_script 生成脚本中的 _EXTRACT_JS 一致。
# 裸标签选择器（如 h3、无 # . [ 空格）在存在 <main> / [role=main] 时只在该区域内匹配，
# 避免 Yahoo 等站顶栏 mega-menu 的 h3 先于正文被 slice(0,n) 取走。
_EXTRACT_JS_MIN = (
    '([s,n])=>{const r=document.querySelector("main")||document.querySelector(\'[role="main"]\');'
    'const bare=/^[a-zA-Z][a-zA-Z0-9-]*$/.test(s)&&s.indexOf("#")<0&&s.indexOf(".")<0&&'
    's.indexOf("[")<0&&s.indexOf(" ")<0;'
    'const sc=bare&&r?r:document;return Array.from(sc.querySelectorAll(s)).slice(0,n)'
    '.map(e=>(e.textContent||"").replace(/\\s+/g," ").trim()).filter(Boolean)}'
)


# ── Code generation helpers ──────────────────────────────────────────────────

def _step_code(step_n: int, context: str, body: list[str]) -> str:
    """Wrap body lines in a try/except block inside async def run()."""
    ind = "            "  # 12 spaces — sits inside: async with → browser → try
    lines = [
        f"{ind}# ── 步骤 {step_n}：{context}",
        f"{ind}try:",
    ]
    for b in body:
        lines.append(f"{ind}    {b}")
    lines += [
        f"{ind}except Exception:",
        f'{ind}    await page.screenshot(path="step_{step_n}_error.png")',
        f"{ind}    raise",
    ]
    return "\n".join(lines)


# params / headers 字符串值填 __ENV:YOUR_ENV_VAR_NAME__，录制与回放均从环境变量读取
_ENV_PLACEHOLDER_RE = re.compile(r"^__ENV:([A-Za-z_][A-Za-z0-9_]*)__$")


def _resolve_placeholders_for_record(v, data: dict):
    """将占位符替换为环境变量或 data['env'] 中的值（仅录制时用于真实 HTTP 请求）。"""
    if not isinstance(v, str):
        return v
    env_fallback = data.get("env") or {}
    m = _ENV_PLACEHOLDER_RE.match(v)
    if m:
        name = m.group(1)
        return os.environ.get(name, env_fallback.get(name, ""))
    return v


def _params_for_record(data: dict) -> dict:
    params = dict(data.get("params") or {})
    for k, v in list(params.items()):
        params[k] = _resolve_placeholders_for_record(v, data)
    return params


def _headers_for_record(data: dict):
    h = data.get("headers")
    if not h:
        return None
    return {k: _resolve_placeholders_for_record(v, data) for k, v in h.items()}


def _build_api_url_for_record(data: dict) -> str:
    if data.get("base_url") is not None and data.get("params") is not None:
        base = str(data["base_url"]).rstrip("?")
        q = urllib.parse.urlencode(_params_for_record(data))
        return f"{base}?{q}"
    url = data.get("url")
    if not url:
        raise ValueError("api_call 需要 url，或 base_url + params / api_call requires 'url' or 'base_url' + 'params'")
    return str(url)


from typing import Optional
def _codegen_env_value(v, env_values: Optional[dict] = None) -> str:
    """params/headers 里单个值对应的 Python 表达式。

    若 v 是 __ENV:VAR__ 占位符：
    - 若 env_values 中有该变量的真实值（用户在录制时已提供），直接写入脚本，免去 export。
    - 否则生成 os.environ.get(...)，运行时从环境变量读取。
    """
    if not isinstance(v, str):
        return repr(v)
    m = _ENV_PLACEHOLDER_RE.match(v)
    if m:
        name = m.group(1)
        if env_values and env_values.get(name):
            # User already supplied the real value — embed it directly so the
            # generated script runs without any `export` setup.
            return repr(env_values[name])
        # Use double quotes consistently so _build_final_script regex can find them.
        return f'os.environ.get("{name}", "")'
    return repr(v)


def _api_codegen_body(context: str, data: dict) -> list[str]:
    """生成 run() 内 httpx 调用代码行（不含 try/except 外壳）。"""
    method = (data.get("method") or "GET").upper()
    save_to = data.get("save_response_to")
    headers = data.get("headers")
    env_values: dict = data.get("env") or {}
    # 自动检测含下划线的主机名 → 生成 verify=False
    _url_for_check = data.get("url") or data.get("base_url") or ""
    _host_for_check = urllib.parse.urlparse(_url_for_check).hostname or ""
    _auto_no_verify = "_" in _host_for_check
    verify_ssl = data.get("verify_ssl", not _auto_no_verify)
    verify_kw = "" if verify_ssl else ", verify=False"

    lines: list[str] = []
    if data.get("base_url") is not None and data.get("params") is not None:
        parts: list[str] = []
        for k, v in data["params"].items():
            parts.append(f"{repr(k)}: {_codegen_env_value(v, env_values)}")
        lines.append("_params = {" + ", ".join(parts) + "}")
        lines.append(
            f'_api_url = {repr(str(data["base_url"]).rstrip("?"))} + "?" + urllib.parse.urlencode(_params)'
        )
    else:
        url = data.get("url")
        lines.append(f"_api_url = {repr(url)}")

    hdr_kw = ""
    if headers:
        hparts = [f"{repr(k)}: {_codegen_env_value(v, env_values)}" for k, v in headers.items()]
        lines.insert(0, "_api_headers = {" + ", ".join(hparts) + "}")
        hdr_kw = ", headers=_api_headers"

    client_kw = f'timeout=CONFIG["api_timeout"]{verify_kw}'
    if method == "GET":
        lines.append(f"async with httpx.AsyncClient({client_kw}) as _hc:")
        lines.append(f"    _r = await _hc.get(_api_url{hdr_kw})")
        lines.append("    _r.raise_for_status()")
    elif method == "POST":
        body = data.get("body")
        lines.append(f"async with httpx.AsyncClient({client_kw}) as _hc:")
        if isinstance(body, dict):
            if hdr_kw:
                lines.append(f"    _r = await _hc.post(_api_url, json={repr(body)}, headers=_api_headers)")
            else:
                lines.append(f"    _r = await _hc.post(_api_url, json={repr(body)})")
        else:
            b = body if body is not None else ""
            if hdr_kw:
                lines.append(f"    _r = await _hc.post(_api_url, content={repr(b)}, headers=_api_headers)")
            else:
                lines.append(f"    _r = await _hc.post(_api_url, content={repr(b)})")
        lines.append("    _r.raise_for_status()")
    else:
        lines.append(f"async with httpx.AsyncClient({client_kw}) as _hc:")
        lines.append(
            f"    _r = await _hc.request({repr(method)}, _api_url{hdr_kw})"
        )
        lines.append("    _r.raise_for_status()")

    if save_to:
        lines.append(f'(CONFIG["output_dir"] / {repr(save_to)}).write_text(_r.text, encoding="utf-8")')
        lines.append(f'print("API 响应已写入", CONFIG["output_dir"] / {repr(save_to)})')
    else:
        lines.append('print("API 响应长度:", len(_r.text))')

    return lines


def _resolve_file(file_str: str, base_dir: Path) -> Path:
    """若 file_str 是绝对路径直接返回；否则相对 base_dir 解析。"""
    p = Path(file_str)
    return p if p.is_absolute() else base_dir / file_str


def _excel_rows_from_json(spec: dict, base_dir: Path) -> list[list]:
    """展平 JSON 文件中的嵌套数组，返回二维行列表。

    spec 格式（任选其一）：
      平铺：{"file":"x.json","outer_key":"items","fields":["f1","f2"]}
      嵌套：{"file":"x.json","outer_key":"batches","inner_key":"lines",
             "fields":["f1","f2"],"parent_fields":["batch_id"]}
    file 可为绝对路径，也可为相对于 base_dir 的文件名。
    """
    import json as _json
    fpath = _resolve_file(spec["file"], base_dir)
    if not fpath.exists():
        return []
    data = _json.loads(fpath.read_text(encoding="utf-8"))
    outer_key = spec.get("outer_key", "")
    inner_key = spec.get("inner_key", "")
    fields = spec.get("fields") or []
    parent_fields = spec.get("parent_fields") or []
    outer_list = data.get(outer_key, []) if outer_key else (data if isinstance(data, list) else [])
    rows: list[list] = []
    if inner_key:
        for outer_item in outer_list:
            pvals = [outer_item.get(pf) for pf in parent_fields]
            for inner_item in outer_item.get(inner_key, []):
                rows.append([inner_item.get(f) for f in fields] + pvals)
    else:
        for item in outer_list:
            rows.append([item.get(f) for f in fields] + [item.get(pf) for pf in parent_fields])
    return rows


def _excel_rows_from_excel(spec: dict, base_dir: Path) -> list[list]:
    """从另一个 xlsx 文件的指定 sheet 读取数据行（不含表头首行）。

    spec 格式：{"file":"发票导入_本周.xlsx","sheet":"发票侧","skip_header":true}
    file 可为绝对路径，也可为相对于 base_dir 的文件名。
    skip_header 默认 true，跳过第一行表头。
    """
    from openpyxl import load_workbook as _lw
    fpath = _resolve_file(spec["file"], base_dir)
    if not fpath.exists():
        return []
    wb = _lw(str(fpath), read_only=True, data_only=True)
    sheet_name = spec.get("sheet", "")
    ws = wb[sheet_name] if sheet_name and sheet_name in wb.sheetnames else wb.active
    skip = bool(spec.get("skip_header", True))
    rows: list[list] = []
    for i, row in enumerate(ws.iter_rows(values_only=True)):
        if i == 0 and skip:
            continue
        rows.append(list(row))
    wb.close()
    return rows


from typing import Optional
def _excel_write_run(data: dict) -> Optional[str]:
    """在录制时写入桌面 xlsx；失败返回错误信息。"""
    try:
        from openpyxl import Workbook, load_workbook
        from openpyxl.utils import get_column_letter
    except ImportError:
        return "缺少 openpyxl：请执行 python3 rpa_manager.py deps-install B 或 pip install openpyxl / Missing openpyxl: run 'python3 rpa_manager.py deps-install B' or 'pip install openpyxl'"

    rel = (data.get("path") or data.get("value") or "").strip()
    sheet = (data.get("sheet") or "").strip()
    if not rel or not sheet:
        return "excel_write 需要 path（或 value）与 sheet / excel_write requires 'path' (or 'value') and 'sheet'"

    output_dir = Path.home() / "Desktop"  # recorder always writes to Desktop
    headers = data.get("headers") or []
    # Dynamic row sources (take precedence over static "rows")
    if data.get("rows_from_json"):
        rows = _excel_rows_from_json(data["rows_from_json"], output_dir)
    elif data.get("rows_from_excel"):
        rows = _excel_rows_from_excel(data["rows_from_excel"], output_dir)
    else:
        rows = data.get("rows") or []
    freeze = (data.get("freeze_panes") or "").strip() or None
    hidden_cols = data.get("hidden_columns") or []
    replace_sheet = bool(data.get("replace_sheet", True))

    path = desktop / rel
    path.parent.mkdir(parents=True, exist_ok=True)

    if path.exists():
        wb = load_workbook(path)
    else:
        wb = Workbook()
        wb.remove(wb.active)

    if sheet in wb.sheetnames:
        if replace_sheet:
            del wb[sheet]
            ws = wb.create_sheet(sheet)
            ridx = 1
            for c, h in enumerate(headers, 1):
                ws.cell(row=ridx, column=c, value=h)
            if headers:
                ridx += 1
            for row in rows:
                for c, v in enumerate(row, 1):
                    ws.cell(row=ridx, column=c, value=v)
                ridx += 1
            if freeze:
                ws.freeze_panes = freeze
            for ci in hidden_cols:
                if isinstance(ci, int) and ci > 0:
                    ws.column_dimensions[get_column_letter(ci)].hidden = True
        else:
            ws = wb[sheet]
            nr = ws.max_row + 1
            for row in rows:
                for c, v in enumerate(row, 1):
                    ws.cell(row=nr, column=c, value=v)
                nr += 1
            if freeze:
                ws.freeze_panes = freeze
            for ci in hidden_cols:
                if isinstance(ci, int) and ci > 0:
                    ws.column_dimensions[get_column_letter(ci)].hidden = True
    else:
        ws = wb.create_sheet(sheet)
        ridx = 1
        for c, h in enumerate(headers, 1):
            ws.cell(row=ridx, column=c, value=h)
        if headers:
            ridx += 1
        for row in rows:
            for c, v in enumerate(row, 1):
                ws.cell(row=ridx, column=c, value=v)
            ridx += 1
        if freeze:
            ws.freeze_panes = freeze
        for ci in hidden_cols:
            if isinstance(ci, int) and ci > 0:
                ws.column_dimensions[get_column_letter(ci)].hidden = True

    wb.save(path)
    print(f"[recorder] excel_write → {path}", flush=True)
    return None


def _excel_write_codegen_lines(data: dict) -> list[str]:
    """与 _excel_write_run 同语义；生成 run() 内异步函数中的同步 openpyxl 代码。"""
    rel = (data.get("path") or data.get("value") or "").strip()
    sheet = (data.get("sheet") or "").strip()
    headers = data.get("headers") or []
    freeze = (data.get("freeze_panes") or "").strip() or None
    hidden_cols = data.get("hidden_columns") or []
    replace_sheet = bool(data.get("replace_sheet", True))

    h_repr = repr(headers)
    rel_repr = repr(rel)
    sheet_repr = repr(sheet)
    freeze_py = repr(freeze) if freeze else "None"
    hid_repr = repr(hidden_cols)
    rep_py = "True" if replace_sheet else "False"

    # --- Dynamic row-loading preamble ---
    row_preamble: list[str] = []
    rfj = data.get("rows_from_json")
    rfe = data.get("rows_from_excel")
    def _file_path_expr(file_str: str) -> str:
        """生成 Path 表达式：绝对路径直接 Path(...)，否则 CONFIG["output_dir"] / ..."""
        if Path(file_str).is_absolute():
            return f"Path({repr(file_str)})"
        return f'CONFIG["output_dir"] / {repr(file_str)}'

    if rfj:
        outer_key = rfj.get("outer_key", "")
        inner_key = rfj.get("inner_key", "")
        fields = rfj.get("fields") or []
        parent_fields = rfj.get("parent_fields") or []
        row_preamble += [
            f"_rfj_path = {_file_path_expr(rfj['file'])}",
            "import json as _json",
            "_rfj_data = _json.loads(_rfj_path.read_text(encoding='utf-8'))",
            "_rows = []",
        ]
        if inner_key:
            row_preamble += [
                f"for _outer in _rfj_data.get({repr(outer_key)}, []):",
                f"    _pvals = [_outer.get(_pf) for _pf in {repr(parent_fields)}]",
                f"    for _inner in _outer.get({repr(inner_key)}, []):",
                f"        _rows.append([_inner.get(_f) for _f in {repr(fields)}] + _pvals)",
            ]
        else:
            _src = f"_rfj_data.get({repr(outer_key)}, [])" if outer_key else "_rfj_data"
            row_preamble += [
                f"for _item in {_src}:",
                f"    _rows.append([_item.get(_f) for _f in {repr(fields)}]"
                + (f" + [_item.get(_pf) for _pf in {repr(parent_fields)}]" if parent_fields else "") + ")",
            ]
    elif rfe:
        src_sheet = repr(rfe.get("sheet", ""))
        skip = repr(bool(rfe.get("skip_header", True)))
        row_preamble += [
            f"_rfe_wb = load_workbook(str({_file_path_expr(rfe['file'])}), read_only=True, data_only=True)",
            f"_rfe_sn = {src_sheet}",
            "_rfe_ws = _rfe_wb[_rfe_sn] if _rfe_sn and _rfe_sn in _rfe_wb.sheetnames else _rfe_wb.active",
            "_rows = []",
            f"for _ri, _rrow in enumerate(_rfe_ws.iter_rows(values_only=True)):",
            f"    if _ri == 0 and {skip}: continue",
            "    _rows.append(list(_rrow))",
            "_rfe_wb.close()",
        ]
    else:
        rows = data.get("rows") or []
        row_preamble = [f"_rows = {repr(rows)}"]

    def _hide(indent: str) -> list[str]:
        return [
            indent + f"for _ci in {hid_repr}:",
            indent + "    if isinstance(_ci, int) and _ci > 0:",
            indent + "        _ws.column_dimensions[get_column_letter(_ci)].hidden = True",
        ]

    lines: list[str] = row_preamble + [
        "_xp = CONFIG[\"output_dir\"] / " + rel_repr,
        "_xp.parent.mkdir(parents=True, exist_ok=True)",
        "if _xp.exists():",
        "    _wb = load_workbook(_xp)",
        "else:",
        "    _wb = Workbook()",
        "    _wb.remove(_wb.active)",
        f"_sh = {sheet_repr}",
        f"_hdrs = {h_repr}",
        f"_replace = {rep_py}",
        "if _sh in _wb.sheetnames:",
        "    if _replace:",
        "        del _wb[_sh]",
        "        _ws = _wb.create_sheet(_sh)",
        "        _ridx = 1",
        "        for _c, _h in enumerate(_hdrs, 1):",
        "            _ws.cell(row=_ridx, column=_c, value=_h)",
        "        if _hdrs:",
        "            _ridx += 1",
        "        for _row in _rows:",
        "            for _c, _v in enumerate(_row, 1):",
        "                _ws.cell(row=_ridx, column=_c, value=_v)",
        "            _ridx += 1",
    ]
    if freeze:
        lines.append("        _ws.freeze_panes = " + freeze_py)
    lines.extend(_hide("        "))
    lines += [
        "    else:",
        "        _ws = _wb[_sh]",
        "        _nr = _ws.max_row + 1",
        "        for _row in _rows:",
        "            for _c, _v in enumerate(_row, 1):",
        "                _ws.cell(row=_nr, column=_c, value=_v)",
        "            _nr += 1",
    ]
    if freeze:
        lines.append("        _ws.freeze_panes = " + freeze_py)
    lines.extend(_hide("        "))
    lines += [
        "else:",
        "    _ws = _wb.create_sheet(_sh)",
        "    _ridx = 1",
        "    for _c, _h in enumerate(_hdrs, 1):",
        "        _ws.cell(row=_ridx, column=_c, value=_h)",
        "    if _hdrs:",
        "        _ridx += 1",
        "    for _row in _rows:",
        "        for _c, _v in enumerate(_row, 1):",
        "            _ws.cell(row=_ridx, column=_c, value=_v)",
        "        _ridx += 1",
    ]
    if freeze:
        lines.append("    _ws.freeze_panes = " + freeze_py)
    lines.extend(_hide("    "))
    lines += ["_wb.save(_xp)", 'print("excel_write →", _xp)']
    return lines


def _python_snippet_run(code: str) -> Optional[str]:
    """在录制时执行 python_snippet 代码，验证依赖和逻辑正确性。

    构建与生成脚本 run() 函数体完全同构的执行命名空间：
      - CONFIG["output_dir"] = ~/Desktop（录制期间的输出目录）
      - 标准库：Path, json, os, datetime, re
      - openpyxl: Workbook, load_workbook, get_column_letter（若已安装）
      - python-docx: Document（若已安装）
      - page = None（浏览器页面对象；非浏览器步骤不可用）

    返回错误字符串；None 表示成功。
    """
    import traceback as _tb

    ns: dict = {
        "Path": Path,
        "CONFIG": {
            "output_dir": Path.home() / "Desktop",
            "task_name": "preview",
        },
        "json": __import__("json"),
        "os": __import__("os"),
        "re": __import__("re"),
        "datetime": __import__("datetime"),
        "page": None,
    }

    # openpyxl
    _missing_deps: list[str] = []
    try:
        from openpyxl import Workbook, load_workbook
        from openpyxl.utils import get_column_letter
        ns.update({"Workbook": Workbook, "load_workbook": load_workbook,
                   "get_column_letter": get_column_letter})
    except ImportError:
        _missing_deps.append("openpyxl")

    # python-docx
    try:
        from docx import Document
        ns["Document"] = Document
    except ImportError:
        _missing_deps.append("python-docx")

    # Pre-check: if code references openpyxl/docx symbols but dep is missing, fail fast
    needs_openpyxl = any(sym in code for sym in ("load_workbook", "Workbook", "openpyxl", "get_column_letter"))
    needs_docx = any(sym in code for sym in ("Document", "from docx", "python_docx"))
    missing_required = []
    if needs_openpyxl and "openpyxl" in _missing_deps:
        missing_required.append("openpyxl")
    if needs_docx and "python-docx" in _missing_deps:
        missing_required.append("python-docx")
    if missing_required:
        pkgs = " ".join(missing_required)
        cap = "B" if "openpyxl" in missing_required else "C"
        return (f"python_snippet 缺少依赖：{pkgs}。请先执行：python3 rpa_manager.py deps-install {cap}"
                f" / python_snippet missing deps: {pkgs}. Run: python3 rpa_manager.py deps-install {cap}")

    # Compile (syntax check)
    try:
        compiled = compile(code, "<python_snippet>", "exec")
    except SyntaxError as e:
        return f"python_snippet 语法错误 / syntax error: {e}"

    # Execute
    try:
        exec(compiled, ns)  # noqa: S102
        print("[recorder] python_snippet 验证通过 / validation passed ✓", flush=True)
        return None
    except ImportError as e:
        mod = str(e).split("'")[1] if "'" in str(e) else str(e)
        return (f"python_snippet ImportError：{e}。请确认 {mod} 已安装（deps-install 或 pip install {mod}）"
                f" / ImportError: {e}. Make sure '{mod}' is installed (deps-install or pip install {mod}).")
    except FileNotFoundError as e:
        return (f"python_snippet 文件未找到：{e}。请确认前序步骤（api_call / excel_write）已成功执行并在桌面生成了该文件"
                f" / FileNotFoundError: {e}. Make sure the preceding api_call / excel_write step ran successfully and the file exists on the Desktop.")
    except Exception as e:
        tb = _tb.format_exc(limit=5)
        return f"python_snippet 执行失败 / execution failed: {type(e).__name__}: {e}\n{tb}"


def _word_write_run(data: dict) -> Optional[str]:
    try:
        from docx import Document
    except ImportError:
        return "缺少 python-docx：请执行 python3 rpa_manager.py deps-install C 或 pip install python-docx / Missing python-docx: run 'python3 rpa_manager.py deps-install C' or 'pip install python-docx'"

    rel = (data.get("path") or data.get("value") or "").strip()
    if not rel:
        return "word_write 需要 path（或 value）/ word_write requires 'path' (or 'value')"

    paragraphs = data.get("paragraphs") or []
    if not isinstance(paragraphs, list):
        return "paragraphs 须为字符串数组 / 'paragraphs' must be an array of strings"

    table_def = data.get("table")  # optional: {"headers": [...], "rows": [[...]]}

    mode = (data.get("mode") or "new").lower()
    path = Path.home() / "Desktop" / rel
    path.parent.mkdir(parents=True, exist_ok=True)

    try:
        if mode == "append" and path.exists():
            doc = Document(str(path))
        else:
            doc = Document()
        for p in paragraphs:
            doc.add_paragraph(str(p))
        if table_def and isinstance(table_def, dict):
            headers = table_def.get("headers") or []
            rows = table_def.get("rows") or []
            col_count = len(headers) or (len(rows[0]) if rows else 0)
            if col_count:
                tbl = doc.add_table(rows=1 + len(rows), cols=col_count)
                tbl.style = "Table Grid"
                hdr_cells = tbl.rows[0].cells
                for i, h in enumerate(headers):
                    hdr_cells[i].text = str(h)
                for r_idx, row in enumerate(rows):
                    cells = tbl.rows[r_idx + 1].cells
                    for c_idx, val in enumerate(row):
                        cells[c_idx].text = str(val)
        doc.save(str(path))
        print(f"[recorder] word_write → {path}", flush=True)
    except Exception as exc:
        return str(exc)
    return None


def _word_write_codegen_lines(data: dict) -> list[str]:
    rel = (data.get("path") or data.get("value") or "").strip()
    paragraphs = data.get("paragraphs") or []
    table_def = data.get("table")
    mode = (data.get("mode") or "new").lower()
    rel_repr = repr(rel)
    par_repr = repr(paragraphs)
    mode_repr = repr(mode)

    lines = [
        "_wp = CONFIG[\"output_dir\"] / " + rel_repr,
        "_wp.parent.mkdir(parents=True, exist_ok=True)",
        f"_wparas = {par_repr}",
        f"_wmode = {mode_repr}",
        'if _wmode == "append" and _wp.exists():',
        "    _doc = Document(str(_wp))",
        "else:",
        "    _doc = Document()",
        "for _p in _wparas:",
        "    _doc.add_paragraph(str(_p))",
    ]

    if table_def and isinstance(table_def, dict):
        headers = table_def.get("headers") or []
        rows = table_def.get("rows") or []
        lines += [
            f"_wtbl_headers = {repr(headers)}",
            f"_wtbl_rows = {repr(rows)}",
            "_wtbl_cols = len(_wtbl_headers) or (len(_wtbl_rows[0]) if _wtbl_rows else 0)",
            "if _wtbl_cols:",
            "    _wtbl = _doc.add_table(rows=1 + len(_wtbl_rows), cols=_wtbl_cols)",
            '    _wtbl.style = "Table Grid"',
            "    for _ci, _h in enumerate(_wtbl_headers):",
            "        _wtbl.rows[0].cells[_ci].text = str(_h)",
            "    for _ri, _row in enumerate(_wtbl_rows):",
            "        for _ci, _v in enumerate(_row):",
            "            _wtbl.rows[_ri + 1].cells[_ci].text = str(_v)",
        ]

    lines += [
        "_doc.save(str(_wp))",
        'print("word_write →", _wp)',
    ]
    return lines


def _format_extract_section(field_label: str, lines: list[str]) -> str:
    """Format extracted DOM text: show field name, separator line, then body."""
    name = (field_label or "").strip() or "extract"
    title = f"【字段：{name}】"
    if not lines:
        body = "(no text matched)\n"
    elif len(lines) == 1:
        body = lines[0].strip() + "\n"
    else:
        parts = [f"{i + 1}. {s.strip()}" for i, s in enumerate(lines)]
        body = "\n\n".join(parts) + "\n"
    sep = "─" * 32
    return f"{title}\n{sep}\n{body}\n"


# ── DOM snapshot ─────────────────────────────────────────────────────────────

async def _snapshot(page) -> list[dict]:
    """Return interactive elements with usable CSS selectors for LLM decision.

    Selector priority:
      1. Own  #id / [data-testid] / [aria-label] / tag[name]
      2. Ancestor walk (max 4 levels) — produces  [data-testid="X"] tag
      3. :nth-of-type fallback inside nearest sectioning parent
    Also returns a separate 'sections' list showing content containers
    so the LLM can scope selectors to specific page areas.
    """
    try:
        return await page.evaluate("""() => {
            // ── helpers ──────────────────────────────────────────────────────
            function ownSel(el) {
                if (el.id) return '#' + el.id;
                const tid  = el.getAttribute('data-testid');
                if (tid)  return `[data-testid="${tid}"]`;
                const aria = el.getAttribute('aria-label');
                if (aria) return `[aria-label="${aria}"]`;
                const name = el.getAttribute('name');
                if (name) return `${el.tagName.toLowerCase()}[name="${name}"]`;
                return null;
            }

            function ancestorSel(el) {
                // Walk up max 4 levels; return composite like [data-testid="X"] a h3
                // IMPORTANT: intermediate tags are included to preserve real DOM nesting
                // order, so LLM cannot confuse "a h3" with "h3 a".
                let cur = el.parentElement;
                const midTags = [];   // intermediate tag names (nearest → farthest)
                for (let d = 0; d < 4 && cur; d++, cur = cur.parentElement) {
                    const s = ownSel(cur);
                    if (s) {
                        // midTags are collected nearest-first; reverse so they read
                        // parent→child in CSS order: ancestor > mid1 > mid2 > el
                        const mid = midTags.slice().reverse().join(' ');
                        return mid ? `${s} ${mid} ${el.tagName.toLowerCase()}`
                                   : `${s} ${el.tagName.toLowerCase()}`;
                    }
                    midTags.push(cur.tagName.toLowerCase());
                }
                return null;
            }

            function nthSel(el) {
                // Fallback: find position among siblings of same tag inside nearest section
                const parent = el.parentElement;
                if (!parent) return null;
                const siblings = Array.from(parent.children)
                    .filter(c => c.tagName === el.tagName);
                const idx = siblings.indexOf(el) + 1;
                const ps = ownSel(parent);
                if (ps) return `${ps} > ${el.tagName.toLowerCase()}:nth-of-type(${idx})`;
                return null;
            }

            // ── collect interactive + heading elements ────────────────────
            const TAGS = [
                'input', 'button', 'select', 'textarea', 'a[href]',
                '[role="button"]', '[role="link"]', '[role="searchbox"]',
                '[role="tab"]', 'h1', 'h2', 'h3', 'li'
            ].join(',');

            const visible = Array.from(document.querySelectorAll(TAGS))
                .filter(el => {
                    const r = el.getBoundingClientRect();
                    return r.width > 0 && r.height > 0;
                })
                .slice(0, 100);

            const items = visible.map(el => {
                const tag  = el.tagName.toLowerCase();
                const ph   = el.getAttribute('placeholder') || null;
                const text = (el.textContent || '').replace(/\\s+/g, ' ').trim().slice(0, 70);
                const sel  = ownSel(el) || ancestorSel(el) || nthSel(el);
                return { tag, sel: sel || null, ph, text: text || null };
            }).filter(e => e.sel || e.text);

            // ── collect content sections (for scoped extraction) ──────────
            const SECTION_TAGS = [
                'section', 'article', '[data-testid]', '[id]'
            ].join(',');
            const sections = Array.from(document.querySelectorAll(SECTION_TAGS))
                .filter(el => {
                    const r = el.getBoundingClientRect();
                    return r.width > 100 && r.height > 50;
                })
                .slice(0, 20)
                .map(el => {
                    const s   = ownSel(el);
                    const h   = el.querySelector('h1,h2,h3');
                    const heading = h ? (h.textContent||'').replace(/\\s+/g,' ').trim().slice(0,50) : null;
                    return s ? { sel: s, heading } : null;
                })
                .filter(Boolean)
                .filter((v, i, a) => a.findIndex(x => x.sel === v.sel) === i);

            return { items, sections };
        }""")
    except Exception:
        return {"items": [], "sections": []}


# ── Action executor ──────────────────────────────────────────────────────────

async def _do_action(page, data: dict, step_n: int, shots_dir: Path) -> dict:
    action  = data.get("action", "")
    target  = data.get("target", "")
    value   = data.get("value", "")
    context = data.get("context") or f"步骤 {step_n}"

    code_block        = None
    error             = None
    inspect_children  = None  # dom_inspect: passed through to result JSON for rpa_manager

    # No-browser mode: browser-specific actions are not available.
    # 无浏览器模式：浏览器操作不可用，返回明确错误提示。
    _BROWSER_ACTIONS = {"goto", "fill", "press", "click", "select_option", "extract_text",
                        "wait", "scroll", "scroll_to", "snapshot", "dom_inspect"}
    if page is None and action in _BROWSER_ACTIONS:
        result: dict = {
            "success":    False,
            "error":      (
                f"此任务能力码不含浏览器，不支持 {action!r} 操作。"
                f" / This capability does not include a browser; {action!r} is not available."
            ),
            "code_block": None,
            "screenshot": None,
            "url":        "",
            "snapshot":   [],
            "sections":   [],
        }
        return result

    try:
        if action == "goto":
            await page.goto(target, wait_until="domcontentloaded")
            await page.wait_for_timeout(1500)   # SPA initial render
            code_block = _step_code(step_n, context, [
                f'await page.goto({repr(target)}, wait_until="domcontentloaded")',
                'await page.wait_for_timeout(CONFIG["spa_settle_ms"])',
            ])

        elif action == "fill":
            loc = page.locator(target).first
            await loc.wait_for(state="visible", timeout=20_000)
            await loc.fill(value)
            code_block = _step_code(step_n, context, [
                f'await page.locator({repr(target)}).first.fill({repr(value)})',
            ])

        elif action == "press":
            key = target or "Enter"
            await page.keyboard.press(key)
            await page.wait_for_load_state("domcontentloaded")
            await page.wait_for_timeout(800)   # let SPA finish routing
            code_block = _step_code(step_n, context, [
                f'await page.keyboard.press({repr(key)})',
                'await page.wait_for_load_state("domcontentloaded")',
                'await page.wait_for_timeout(800)',
            ])

        elif action == "click":
            loc = page.locator(target).first
            await loc.wait_for(state="visible", timeout=20_000)
            await loc.click()
            await page.wait_for_load_state("domcontentloaded")
            await page.wait_for_timeout(800)   # let SPA finish routing
            code_block = _step_code(step_n, context, [
                f'await page.locator({repr(target)}).first.click()',
                'await page.wait_for_load_state("domcontentloaded")',
                'await page.wait_for_timeout(800)',
            ])

        elif action == "select_option":
            # Native <select>: target = CSS, value = option value / label / index (see select_by)
            # Playwright fill() does NOT set <select>; use select_option (e.g. Sauce Demo hilo = price high→low)
            loc = page.locator(target).first
            await loc.wait_for(state="visible", timeout=20_000)
            how = (data.get("select_by") or "value").lower().strip()
            if how == "label":
                await loc.select_option(label=value)
                sel_line = f'await page.locator({repr(target)}).first.select_option(label={repr(value)})'
            elif how == "index":
                idx = int(value)
                await loc.select_option(index=idx)
                sel_line = f'await page.locator({repr(target)}).first.select_option(index={idx})'
            else:
                await loc.select_option(value)
                sel_line = f'await page.locator({repr(target)}).first.select_option({repr(value)})'
            await page.wait_for_load_state("domcontentloaded")
            await page.wait_for_timeout(800)
            code_block = _step_code(step_n, context, [
                sel_line,
                'await page.wait_for_load_state("domcontentloaded")',
                'await page.wait_for_timeout(800)',
            ])

        elif action == "extract_text":
            filename = value or "output.txt"
            limit    = int(data.get("limit", 0)) or 0

            # Wait for page to settle (SPA re-renders can cause locator.all() race)
            try:
                await page.wait_for_load_state("domcontentloaded", timeout=5_000)
            except Exception:
                pass

            # Single atomic JS call — immune to mid-flight page re-renders
            limit_n = limit or 9999
            texts = await page.evaluate(_EXTRACT_JS_MIN, [target, limit_n])

            # field / field_name = short label for output; else fall back to context
            field_label = (
                (data.get("field") or data.get("field_name") or context or f"步骤 {step_n}")
                .strip()
                or "extract"
            )
            first_for_name = filename not in _EXTRACT_OUTPUT_FILES
            _EXTRACT_OUTPUT_FILES.add(filename)

            out = Path.home() / "Desktop" / filename
            blob = _format_extract_section(field_label, texts)
            if first_for_name:
                out.write_text(blob, encoding="utf-8")
            else:
                with out.open("a", encoding="utf-8") as f:
                    f.write(blob)
            if texts:
                print(f"[recorder] extracted {len(texts)} items → {out}", flush=True)
            else:
                print(f"[recorder] ⚠️  WARNING: 0 items matched selector {repr(target)}", flush=True)
                print(f"[recorder]    The selector may be wrong or content not yet rendered.", flush=True)
                print(f"[recorder]    Try: dom_inspect on a parent container to see real DOM structure.", flush=True)
                error = (f"⚠️ 提取到 0 条内容。选择器 {repr(target)} 可能不匹配当前页面的真实 DOM 结构。"
                         f"\n建议：先用 dom_inspect 检查父容器的真实子元素结构，再修正选择器。"
                         f"\n / 0 items extracted. Selector {repr(target)} may not match the real DOM structure."
                         f"\nTip: use dom_inspect on a parent container to see actual child elements, then fix the selector.")

            # Generated script: same filename → first step write_text, later steps append
            lim_code = str(limit) if limit else "9999"
            field_lit = repr(
                (data.get("field") or data.get("field_name") or context or f"步骤 {step_n}").strip()
                or "extract"
            )
            common_lines = [
                f'_sel = {repr(target)}',
                f'_lim = {lim_code}',
                'await _wait_for_content(page, _sel)',
                '_texts = await page.evaluate(_EXTRACT_JS, [_sel, _lim])',
                f'_out = CONFIG["output_dir"] / {repr(filename)}',
                f'_field = {field_lit}',
                '_block = _format_extract_section(_field, _texts)',
            ]
            if first_for_name:
                body_lines = common_lines + [
                    '_out.write_text(_block, encoding="utf-8")',
                    'print(f"已提取 {len(_texts)} 条，写入 {_out}（本文件首次写入）")',
                ]
            else:
                body_lines = common_lines + [
                    'with _out.open("a", encoding="utf-8") as _f:',
                    '    _f.write(_block)',
                    'print(f"已提取 {len(_texts)} 条，追加写入 {_out}")',
                ]
            code_block = _step_code(step_n, context, body_lines)

        elif action == "wait":
            ms = int(value) if value else 2000
            await page.wait_for_timeout(ms)
            code_block = _step_code(step_n, context, [
                f'await page.wait_for_timeout({ms})',
            ])

        elif action == "scroll":
            px = int(value) if value else 500
            try:
                await page.wait_for_load_state("domcontentloaded", timeout=10_000)
            except Exception:
                pass
            vp = page.viewport_size
            if vp:
                await page.mouse.move(vp["width"] // 2, vp["height"] // 2)
            else:
                await page.mouse.move(720, 450)
            await page.mouse.wheel(0, float(px))
            await page.wait_for_timeout(600)   # wait for lazy-load trigger
            code_block = _step_code(step_n, context, [
                f"await _scroll_window(page, {px})",
                "await page.wait_for_timeout(600)",
            ])

        elif action == "scroll_to":
            # Scroll a specific element into view — triggers lazy-load for that section
            await page.evaluate(
                """(sel) => {
                    const el = document.querySelector(sel);
                    if (el) el.scrollIntoView({ behavior: 'smooth', block: 'center' });
                }""",
                target,
            )
            await page.wait_for_timeout(1200)  # wait for lazy content to render
            # Use single-quoted JS string in generated code to avoid escape hell
            js_str = "(s)=>{const e=document.querySelector(s);if(e)e.scrollIntoView({block:'center'})}"
            code_block = _step_code(step_n, context, [
                f'await page.evaluate({repr(js_str)}, {repr(target)})',
                'await page.wait_for_timeout(1200)',
            ])

        elif action == "api_call":
            url = _build_api_url_for_record(data)
            method = (data.get("method") or "GET").upper()
            timeout = float(data.get("timeout") or 60.0)
            save_to = data.get("save_response_to")
            headers = _headers_for_record(data)
            body = data.get("body")
            # verify_ssl：用户显式传 false → 跳过；
            # 未传时自动检测：主机名含下划线（SSL 规范不允许）→ 自动关闭校验
            _host = urllib.parse.urlparse(url).hostname or ""
            _has_underscore_host = "_" in _host
            verify_ssl = data.get("verify_ssl", not _has_underscore_host)
            if _has_underscore_host and data.get("verify_ssl") is None:
                print(f"[recorder] ⚠️  主机名含下划线（{_host}），已自动关闭 SSL 校验（verify=False）"
                      f" / hostname contains underscore ({_host}), SSL verification auto-disabled (verify=False)", flush=True)
            try:
                async with httpx.AsyncClient(timeout=timeout, verify=verify_ssl) as _hc:
                    if method == "GET":
                        r = await _hc.get(url, headers=headers)
                    elif method == "POST":
                        if isinstance(body, dict):
                            r = await _hc.post(url, json=body, headers=headers)
                        else:
                            r = await _hc.post(
                                url,
                                content=body if body is not None else "",
                                headers=headers,
                            )
                    else:
                        r = await _hc.request(method, url, headers=headers)
                    r.raise_for_status()
                    _api_text = r.text
            except httpx.ConnectError as _ssl_exc:
                _msg = str(_ssl_exc)
                if "SSL" in _msg or "certificate" in _msg.lower() or "CERTIFICATE" in _msg:
                    raise RuntimeError(
                        f"SSL 证书验证失败（{_msg}）。\n"
                        "若目标服务使用了含下划线的主机名或自签名证书，可在 api_call 步骤中加上 "
                        "\"verify_ssl\": false 跳过校验（仅用于测试环境）。\n"
                        f" / SSL certificate verification failed ({_msg}).\n"
                        "If the host has an underscore in its name or uses a self-signed cert, "
                        "add \"verify_ssl\": false to the api_call step (test environments only)."
                    ) from _ssl_exc
                raise
            if save_to:
                _out = Path.home() / "Desktop" / save_to
                _out.parent.mkdir(parents=True, exist_ok=True)
                _out.write_text(_api_text, encoding="utf-8")
            code_block = _step_code(step_n, context, _api_codegen_body(context, data))

        elif action == "excel_write":
            err_ex = _excel_write_run(data)
            if err_ex:
                error = err_ex
            else:
                code_block = _step_code(step_n, context, _excel_write_codegen_lines(data))

        elif action == "word_write":
            err_wd = _word_write_run(data)
            if err_wd:
                error = err_wd
            else:
                code_block = _step_code(step_n, context, _word_write_codegen_lines(data))

        elif action == "python_snippet":
            # Execute the snippet NOW (same as api_call / excel_write run at record time).
            # This validates dependencies, file existence, and logic before writing code_block.
            raw_code = (data.get("code") or "").rstrip()
            if not raw_code.strip():
                error = "python_snippet 需要非空的 code 字段 / python_snippet requires a non-empty 'code' field"
            else:
                error = _python_snippet_run(raw_code)
                if not error:
                    code_block = _step_code(step_n, context, raw_code.splitlines())

        elif action == "merge_files":
            # Pure file operation: read sources from Desktop, concatenate, write target.
            # Does NOT interact with the browser.
            sources   = data.get("sources") or []  # list of Desktop filenames
            target_fn = data.get("target") or data.get("value") or ""
            separator = data.get("separator", "\n\n")
            if not sources or not target_fn:
                error = "merge_files 需要 sources（列表）和 target（目标文件名）/ merge_files requires 'sources' (list) and 'target' (output filename)"
            else:
                parts: list[str] = []
                for src in sources:
                    p = Path.home() / "Desktop" / src
                    if p.exists():
                        parts.append(p.read_text(encoding="utf-8"))
                    else:
                        print(f"[recorder] ⚠️  merge_files：文件不存在，跳过 / file not found, skipping: {p}", flush=True)
                out_path = Path.home() / "Desktop" / target_fn
                out_path.write_text(separator.join(parts), encoding="utf-8")
                print(f"[recorder] merge_files → {out_path}（{len(parts)}/{len(sources)} 个源文件 / source files merged）", flush=True)
            # Code generation
            sep_repr   = repr(separator)
            srcs_repr  = repr(sources)
            tgt_repr   = repr(target_fn)
            body_lines = [
                f"_merge_sources = {srcs_repr}",
                f"_merge_sep = {sep_repr}",
                "_merge_parts = []",
                "for _src in _merge_sources:",
                '    _p = CONFIG["output_dir"] / _src',
                "    if _p.exists():",
                '        _merge_parts.append(_p.read_text(encoding="utf-8"))',
                '    else:',
                '        print(f"⚠️  merge_files：文件不存在，跳过 {_p}")',
                f'(CONFIG["output_dir"] / {tgt_repr}).write_text(_merge_sep.join(_merge_parts), encoding="utf-8")',
                f'print("已合并到", CONFIG["output_dir"] / {tgt_repr})',
            ]
            code_block = _step_code(step_n, context, body_lines)

        elif action == "snapshot":
            pass  # read-only DOM inspection — NOT logged to script

        elif action == "dom_inspect":
            # Diagnostic: return child structure of a container element.
            # NOT logged to the script — only used to discover real selectors.
            result = await page.evaluate("""(sel) => {
                const el = document.querySelector(sel);
                if (!el) return { found: false, message: 'Element not found: ' + sel };
                const children = Array.from(el.querySelectorAll('*'))
                    .slice(0, 50)
                    .map(c => ({
                        tag:    c.tagName.toLowerCase(),
                        id:     c.id || null,
                        testid: c.getAttribute('data-testid') || null,
                        cls:    Array.from(c.classList).slice(0, 2).join(' ') || null,
                        aria:   c.getAttribute('aria-label') || null,
                        text:   (c.textContent || '').replace(/\\s+/g, ' ').trim().slice(0, 60),
                    }));
                return { found: true, outerTag: el.tagName.toLowerCase(), children };
            }""", target)
            # Attach inspect result directly to action result (not to code_block)
            if isinstance(result, dict) and result.get("found"):
                children = result.get("children", [])
                # Print structured summary for the LLM to read
                lines = [f"[dom_inspect] 容器 {repr(target)} 共 {len(children)} 个子元素 / container {repr(target)} has {len(children)} children:"]
                for c in children[:30]:
                    sel_hint = (
                        f"#{c['id']}" if c['id']
                        else f"[data-testid=\"{c['testid']}\"]" if c['testid']
                        else f"[aria-label=\"{c['aria']}\"]" if c['aria']
                        else f".{c['cls'].split()[0]}" if c['cls']
                        else c['tag']
                    )
                    print(f"  {sel_hint}  「{c['text'][:50]}」", flush=True)
                # Return children in result for rpa_manager to display
                code_block = None  # not logged
                inspect_children = children
            else:
                error = result.get("message", "dom_inspect failed") if isinstance(result, dict) else "dom_inspect failed"

        else:
            error = f"未知 action / unknown action: {action!r}"

    except Exception as exc:
        error = str(exc)

    # Screenshot after every action (proof + visual feedback for user)
    shot_path: "Path | None" = None
    snap: list = []
    sections: list = []
    url = ""

    if page is not None:
        ts        = datetime.now().strftime("%H%M%S")
        label     = "snapshot" if action == "snapshot" else f"step_{step_n:02d}"
        shot_path = shots_dir / f"{label}_{ts}.png"
        try:
            await page.screenshot(path=str(shot_path), full_page=False)
        except Exception:
            shot_path = None

        # Always return current DOM snapshot so LLM can choose next selector
        raw_snap = await _snapshot(page)
        # _snapshot now returns {"items": [...], "sections": [...]}
        if isinstance(raw_snap, dict):
            snap     = raw_snap.get("items", [])
            sections = raw_snap.get("sections", [])
        else:
            snap = raw_snap  # backward compat if something went wrong

        url = page.url

    out = {
        "success":    error is None,
        "error":      error,
        "code_block": code_block,
        "screenshot": str(shot_path) if shot_path else None,
        "url":        url,
        "snapshot":   snap,
        "sections":   sections,
    }
    if inspect_children is not None:
        out["_inspect_children"] = inspect_children
    return out


# ── Script builder ───────────────────────────────────────────────────────────

def _build_final_script(
    task_name: str,
    code_blocks: list[str],
    *,
    use_openpyxl: bool = False,
    use_docx: bool = False,
    cookies_file: str = "",
) -> str:
    ts    = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    steps = "\n\n".join(code_blocks) if code_blocks else "            pass  # 无录制步骤"
    fmt_src = inspect.getsource(_format_extract_section)
    extract_js_repr = repr(_EXTRACT_JS_MIN)

    # Collect env vars used by __ENV:VAR__ placeholders in all api_call steps.
    # They appear in code blocks as:  os.environ.get("VAR_NAME", "")
    _env_var_re = re.compile(r'os\.environ\.get\("([A-Za-z_][A-Za-z0-9_]*)",\s*""\)')
    env_vars = sorted(set(_env_var_re.findall(steps)))

    # Replace inline os.environ.get(...) with CONFIG["VAR"] so keys are centralised.
    for _v in env_vars:
        steps = steps.replace(f'os.environ.get("{_v}", "")', f'CONFIG["{_v}"]')

    # Build CONFIG entries for env vars (one line each, with export hint).
    if env_vars:
        env_config_lines = "\n    # --- API 密钥（从环境变量读取；个人使用时也可直接在此填写）"
        for _v in env_vars:
            env_config_lines += f'\n    "{_v}": os.environ.get("{_v}", ""),'
        env_config_lines += "\n"
    else:
        env_config_lines = ""

    # Startup check: fail fast with helpful message if any required key is empty.
    if env_vars:
        env_list_repr = repr(env_vars)
        startup_check = f"""\
_missing = [v for v in {env_list_repr} if not CONFIG.get(v)]
if _missing:
    print("\\n❌  以下 API 密钥未配置，脚本无法运行 / The following API keys are not set, cannot run:")
    for _v in _missing:
        print(f"   {{_v}} 为空 / is empty  →  export {{_v}}='your-key'")
        print(f"   （或在脚本 CONFIG 字典里直接填写 / or set it directly in the CONFIG dict above）")
    raise SystemExit(1)

"""
    else:
        startup_check = ""

    office_imports = ""
    if use_openpyxl:
        office_imports += (
            "\nfrom openpyxl import Workbook, load_workbook\n"
            "from openpyxl.utils import get_column_letter\n"
        )
    if use_docx:
        office_imports += "\nfrom docx import Document\n"

    cookies_path_line = (
        f'\n    # 已保存的登录 Cookie 路径（由 #rpa-login 生成；留空则不注入）\n'
        f'    "cookies_path":  {repr(cookies_file)},'
        if cookies_file else
        '\n    # 若有已保存的登录 Cookie，填写路径（由 rpa_manager login-start 生成）\n'
        '    "cookies_path":  "",'
    )

    return f"""\
# pip install playwright httpx openpyxl python-docx && playwright install chromium
# 任务：{task_name}
# 录制时间：{ts}
# 由 OpenClaw RPA Recorder（headed 真实录制）生成 — 可脱离 OpenClaw 独立运行

import asyncio
import os
import urllib.parse
from pathlib import Path

import httpx{office_imports}
from playwright.async_api import async_playwright, TimeoutError as PlaywrightTimeout

CONFIG = {{
    "output_dir":    Path.home() / "Desktop",
    "headless":      False,
    "timeout":       60_000,
    "slow_mo":       300,
    # 导航后等待 SPA 内容渲染的额外时间（重型 SPA 如 Yahoo Finance 需要 1-2 秒）
    "spa_settle_ms": 1_500,
    # extract_text 等待目标元素出现的超时（毫秒）
    "content_wait":  15_000,
    # httpx 调用外部 API 的超时（秒）
    "api_timeout":   60.0,{cookies_path_line}{env_config_lines}}}

{startup_check}_UA = (
    "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
    "AppleWebKit/537.36 (KHTML, like Gecko) "
    "Chrome/124.0.0.0 Safari/537.36"
)

{fmt_src}

_EXTRACT_JS = {extract_js_repr}


async def _wait_for_content(page, selector: str) -> None:
    \"\"\"等待 selector 对应的元素出现在 DOM 中（容错：超时也继续）。\"\"\"
    try:
        await page.wait_for_selector(selector, timeout=CONFIG["content_wait"])
    except Exception:
        pass  # 元素未出现也继续，evaluate 会返回空列表


async def _scroll_window(page, dy: int) -> None:
    \"\"\"窗口滚动：导航后若再用 evaluate(scrollBy)，易因执行上下文销毁报错；用 mouse.wheel 并在滚动前等待页面稳定。\"\"\"
    try:
        await page.wait_for_load_state("domcontentloaded", timeout=10_000)
    except Exception:
        pass
    vp = page.viewport_size
    if vp:
        await page.mouse.move(vp["width"] // 2, vp["height"] // 2)
    else:
        await page.mouse.move(720, 450)
    await page.mouse.wheel(0, float(dy))


async def run():
    async with async_playwright() as p:
        browser = await p.chromium.launch(
            headless=CONFIG["headless"],
            args=["--no-sandbox", "--disable-blink-features=AutomationControlled"],
            slow_mo=CONFIG["slow_mo"],
        )
        context = await browser.new_context(
            user_agent=_UA,
            viewport={{"width": 1440, "height": 900}},
            locale="en-US",
            timezone_id="America/New_York",
            extra_http_headers={{"Accept-Language": "en-US,en;q=0.9"}},
        )
        await context.add_init_script(
            "Object.defineProperty(navigator, 'webdriver', {{get: () => undefined}})"
        )
        # 若配置了 cookies_path，注入 Cookie 模拟已登录态（由 rpa_manager login-start 生成）
        import json as _json
        _cp = CONFIG.get("cookies_path", "")
        if _cp and Path(_cp).exists():
            await context.add_cookies(_json.loads(Path(_cp).read_text()))
        page = await context.new_page()
        page.set_default_timeout(CONFIG["timeout"])

        try:
{steps}

        except PlaywrightTimeout as e:
            await page.screenshot(path="error_timeout.png")
            raise RuntimeError(f"超时：{{e}}") from e
        except Exception:
            await page.screenshot(path="error_unexpected.png")
            raise
        finally:
            await browser.close()


if __name__ == "__main__":
    asyncio.run(run())
"""


# ── Login capture mode / 登录捕获模式 ────────────────────────────────────────

async def login_capture_main():
    """登录捕获模式：打开浏览器到登录页，等待用户完成登录，再导出 Cookie 保存到文件。
    Login capture mode: open browser to login page, wait for user to finish login,
    then export cookies to a file for later injection."""
    SESSION_DIR.mkdir(parents=True, exist_ok=True)
    pid_path = SESSION_DIR / "server.pid"
    pid_path.write_text(str(os.getpid()))

    task_data         = json.loads((SESSION_DIR / "task.json").read_text())
    login_url         = task_data["login_url"]
    cookies_output    = task_data.get("cookies_output", str(SESSION_DIR / "cookies.json"))
    cookies_meta_out  = task_data.get("cookies_meta_output", str(SESSION_DIR / "cookies_meta.json"))

    from playwright.async_api import async_playwright

    async with async_playwright() as p:
        browser = await p.chromium.launch(
            headless=False,
            args=["--no-sandbox", "--disable-blink-features=AutomationControlled"],
            slow_mo=0,
        )
        ctx = await browser.new_context(
            user_agent=_UA,
            viewport={"width": 1440, "height": 900},
            locale="zh-CN",
            timezone_id="Asia/Shanghai",
            extra_http_headers={"Accept-Language": "zh-CN,zh;q=0.9,en;q=0.8"},
        )
        await ctx.add_init_script(
            "Object.defineProperty(navigator,'webdriver',{get:()=>undefined})"
        )
        page = await ctx.new_page()
        page.set_default_timeout(120_000)

        await page.goto(login_url)

        # 通知 rpa_manager 浏览器已就绪 / Signal rpa_manager that the browser is ready
        (SESSION_DIR / "ready").write_text("1")
        print(f"[login_capture] ready — {login_url}", flush=True)

        # 轮询等待 login_done 指令 / Poll for login_done command from rpa_manager
        cmd_path = SESSION_DIR / "cmd.json"
        last_seq = -1
        while True:
            if cmd_path.exists():
                try:
                    data = json.loads(cmd_path.read_text())
                    seq  = data.get("seq", 0)
                    if seq > last_seq:
                        last_seq = seq
                        if data.get("action") == "login_done":
                            break
                except Exception:
                    pass
            await asyncio.sleep(POLL_INTERVAL)

        # 导出当前 context 的全部 Cookie / Export all cookies from the current browser context
        cookies = await ctx.cookies()
        Path(cookies_output).parent.mkdir(parents=True, exist_ok=True)
        Path(cookies_output).write_text(
            json.dumps(cookies, indent=2, ensure_ascii=False), encoding="utf-8"
        )

        # 生成 meta 摘要 / Build metadata summary (total, session-type count, earliest expiry)
        expiries = [c["expires"] for c in cookies if c.get("expires", -1) > 0]
        earliest = min(expiries) if expiries else None
        meta = {
            "saved_at":       datetime.now().isoformat(timespec="seconds"),
            "total":          len(cookies),
            "session_cookies": sum(1 for c in cookies if c.get("expires", -1) <= 0),
            "earliest_expires": (
                datetime.fromtimestamp(earliest).isoformat(timespec="seconds")
                if earliest else None
            ),
            "hint": "实际有效期以服务端策略为准，可能早于参考时间",
        }
        Path(cookies_meta_out).write_text(
            json.dumps(meta, indent=2, ensure_ascii=False), encoding="utf-8"
        )

        await browser.close()

        # 写 login_done 标记通知 rpa_manager / Write done marker so rpa_manager unblocks
        (SESSION_DIR / "login_done").write_text("1")
    pid_path.unlink(missing_ok=True)
    print(f"[login_capture] done — {len(cookies)} cookies → {cookies_output}", flush=True)


# ── Server main loop ─────────────────────────────────────────────────────────

async def server_main():
    SESSION_DIR.mkdir(parents=True, exist_ok=True)
    pid_path = SESSION_DIR / "server.pid"
    pid_path.write_text(str(os.getpid()))

    _reset_extract_output_tracking()

    task_data = json.loads((SESSION_DIR / "task.json").read_text())
    task_name = task_data["task"]
    shots_dir = SESSION_DIR / "screenshots"
    shots_dir.mkdir(exist_ok=True)

    code_blocks: list[str] = []
    step_n = 0
    use_openpyxl = False
    use_docx = False

    needs_browser = task_data.get("needs_browser", True)

    if not needs_browser:
        # File/API-only mode (B/C/F/N): no Playwright browser needed.
        # 纯文件/API 模式（B/C/F/N）：不启动浏览器。
        (SESSION_DIR / "ready").write_text("1")
        print(f"[recorder] ready (no-browser mode) — task: {task_name}", flush=True)

        last_seq = -1
        cmd_path = SESSION_DIR / "cmd.json"

        while True:
            if cmd_path.exists():
                try:
                    data = json.loads(cmd_path.read_text())
                    seq  = data.get("seq", 0)
                    if seq > last_seq:
                        last_seq = seq
                        action   = data.get("action", "")

                        if action == "shutdown":
                            break

                        if action == "excel_write":
                            use_openpyxl = True
                        if action == "word_write":
                            use_docx = True
                        if action == "python_snippet":
                            snippet_code = (data.get("code") or "")
                            if "load_workbook" in snippet_code or "openpyxl" in snippet_code or "Workbook" in snippet_code:
                                use_openpyxl = True
                            if "Document" in snippet_code or "from docx" in snippet_code:
                                use_docx = True

                        if action != "snapshot":
                            step_n += 1

                        result = await _do_action(None, data, step_n, shots_dir)

                        if result.get("code_block"):
                            code_blocks.append(result["code_block"])

                        (SESSION_DIR / f"result_{seq}.json").write_text(
                            json.dumps(result, ensure_ascii=False, indent=2)
                        )
                except Exception as exc:
                    try:
                        (SESSION_DIR / f"result_{last_seq}.json").write_text(
                            json.dumps(
                                {"success": False, "error": str(exc),
                                 "code_block": None, "snapshot": []},
                                ensure_ascii=False,
                            )
                        )
                    except Exception:
                        pass

            await asyncio.sleep(POLL_INTERVAL)

    else:
        # Browser mode (A/D/E/G): launch headed Chromium.
        # 浏览器模式（A/D/E/G）：启动有界面 Chromium。
        from playwright.async_api import async_playwright

        async with async_playwright() as p:
            browser = await p.chromium.launch(
                headless=False,
                args=["--no-sandbox", "--disable-blink-features=AutomationControlled"],
                slow_mo=200,
            )
            ctx = await browser.new_context(
                user_agent=_UA,
                viewport={"width": 1440, "height": 900},
                locale="en-US",
                timezone_id="America/New_York",
                extra_http_headers={"Accept-Language": "en-US,en;q=0.9"},
            )
            await ctx.add_init_script(
                "Object.defineProperty(navigator,'webdriver',{get:()=>undefined})"
            )

            # 若任务配置了 cookies_file（由 #rpa-autologin 触发），则注入 Cookie 模拟已登录态
            # If task.json has cookies_file (set by #rpa-autologin), inject cookies to simulate logged-in state
            _cookies_file = task_data.get("cookies_file", "")
            if _cookies_file and Path(_cookies_file).exists():
                _raw_cookies = json.loads(Path(_cookies_file).read_text())
                await ctx.add_cookies(_raw_cookies)
                print(f"[recorder] 已注入 {len(_raw_cookies)} 条 Cookie（{_cookies_file}）", flush=True)
            elif _cookies_file:
                print(f"[recorder] ⚠️  cookies_file 不存在，跳过注入：{_cookies_file}", flush=True)

            page = await ctx.new_page()
            page.set_default_timeout(60_000)

            # Signal ready to rpa_manager (it polls for this file)
            (SESSION_DIR / "ready").write_text("1")
            print(f"[recorder] ready — task: {task_name}", flush=True)

            last_seq = -1
            cmd_path = SESSION_DIR / "cmd.json"

            while True:
                if cmd_path.exists():
                    try:
                        data = json.loads(cmd_path.read_text())
                        seq  = data.get("seq", 0)
                        if seq > last_seq:
                            last_seq = seq
                            action   = data.get("action", "")

                            if action == "shutdown":
                                break

                            if action == "excel_write":
                                use_openpyxl = True
                            if action == "word_write":
                                use_docx = True
                            if action == "python_snippet":
                                snippet_code = (data.get("code") or "")
                                if "load_workbook" in snippet_code or "openpyxl" in snippet_code or "Workbook" in snippet_code:
                                    use_openpyxl = True
                                if "Document" in snippet_code or "from docx" in snippet_code:
                                    use_docx = True

                            if action != "snapshot":
                                step_n += 1

                            result = await _do_action(page, data, step_n, shots_dir)

                            if result.get("code_block"):
                                code_blocks.append(result["code_block"])

                            (SESSION_DIR / f"result_{seq}.json").write_text(
                                json.dumps(result, ensure_ascii=False, indent=2)
                            )
                    except Exception as exc:
                        try:
                            (SESSION_DIR / f"result_{last_seq}.json").write_text(
                                json.dumps(
                                    {"success": False, "error": str(exc),
                                     "code_block": None, "snapshot": []},
                                    ensure_ascii=False,
                                )
                            )
                        except Exception:
                            pass

                await asyncio.sleep(POLL_INTERVAL)

            await browser.close()

    # Compile and save final script
    script = _build_final_script(
        task_name,
        code_blocks,
        use_openpyxl=use_openpyxl,
        use_docx=use_docx,
        cookies_file=task_data.get("cookies_file", ""),
    )
    (SESSION_DIR / "script_log.py").write_text(script, encoding="utf-8")
    (SESSION_DIR / "done").write_text("1")
    pid_path.unlink(missing_ok=True)
    print(f"[recorder] done — {len(code_blocks)} steps — script saved.", flush=True)


if __name__ == "__main__":
    _task_file = SESSION_DIR / "task.json"
    _mode = "record"
    if _task_file.exists():
        try:
            _mode = json.loads(_task_file.read_text()).get("mode", "record")
        except Exception:
            pass
    if _mode == "login_capture":
        asyncio.run(login_capture_main())
    else:
        asyncio.run(server_main())
